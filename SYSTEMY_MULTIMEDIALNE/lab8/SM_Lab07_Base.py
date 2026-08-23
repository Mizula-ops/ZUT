import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import soundfile as sf
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO

##########################################
### Settings #############################
##########################################

Only_Tests=False
bit_test=6

DPCM_n=2
DPCM_predictor=np.median

OutputRaportFile = "raport.docx"
OutputFolder="" # place for all your new audio files will be

##########################################
### Data Set #############################
##########################################

AudioDir = r'.' # Address of folder with files (do nor delete `r``)

SingFiles=["SING2/sing_high1.wav","SING2/sing_low1.wav","SING2/sing_medium1.wav"] # list of file names of with Singing Voice

##########################################
### Functions to  ########################
##########################################

def Kwant(data,bit):
    d = (2 ** bit) - 1
    if np.issubdtype(data.dtype, np.floating):
        m = -1
        n = 1
    else:
        m = np.iinfo(data.dtype).min
        n = np.iinfo(data.dtype).max
    # Przygotowanie sygnalu
    DataF = data.astype(float)
    DataF = (DataF - m) / (n - m)
    # Kwantyzacja
    DataF = (np.round(DataF * d)) / d
    # Powrot do orginalnej przestrzeni
    DataF = (DataF * (n - m)) + m
    return DataF.astype(data.dtype)

def A_law_compress(x):
    y=np.zeros(x.shape)
    A=87.6
    absx=np.abs(x)
    idx=absx<(1/A)
    s=np.sign(x)
    y[idx]=s[idx]*(A*absx[idx])/(1+np.log(A))
    y[np.logical_not(idx)]=s[np.logical_not(idx)]*(1+np.log(A*absx[np.logical_not(idx)]))/(1+np.log(A))
    return y



def A_law_decompress(x):
    A=87.6
    y = np.zeros(x.shape)
    absx = np.abs(x)
    idx = absx < (1 / (1+np.log(A)))
    s=np.sign(x)
    y[idx]=s[idx]*(absx[idx]*(1+np.log(A)))/A
    y[np.logical_not(idx)]=s[np.logical_not(idx)]*np.exp(absx[np.logical_not(idx)]*(1+np.log(A))-1)/A

    return y
    
def mu_law_compress(x):
    u=255
    y = np.zeros(x.shape)
    absx = np.abs(x)
    idx = (x >= -1) & (x <= 1)
    s = np.sign(x)
    y[idx] = s[idx] * np.log(1+u*absx[idx])/np.log(1+u)
    y[np.logical_not(idx)] = x[np.logical_not(idx)]
    return y

def mu_law_decompress(x):
    u = 255
    y = np.zeros(x.shape)
    absx = np.abs(x)
    idx = (x >= -1) & (x <= 1)
    s = np.sign(x)
    y[idx] = s[idx] * (((1 + u) ** absx[idx] - 1) / u)
    y[np.logical_not(idx)] = x[np.logical_not(idx)]
    return y

def DPCM_compress(x,bit):
    y=np.zeros(x.shape)
    e=0
    for i in range(0,x.shape[0]):
        y[i]=Kwant(x[i]-e,bit)
        e+=y[i]
    return y

def DPCM_decompress(x):
    y = np.zeros(x.shape)
    e = 0
    for i in range(0, x.shape[0]):
        y[i]=x[i]+e
        e=y[i]
    return y

def DPCM_compress_pred(x,bit,n,predictor=np.mean): 
    y=np.zeros(x.shape)
    xp=np.zeros(x.shape)
    e=0
    for i in range(0,x.shape[0]):
        y[i]=Kwant(x[i]-e,bit)
        xp[i]=y[i]+e
        idx=(np.arange(i-n,i,1,dtype=int)+1)
        idx=np.delete(idx,idx<0)
        e=predictor(xp[idx])
    return y

def DPCM_decompress_pred(x,n,predictor=np.mean):
    y = np.zeros(x.shape)
    e = 0
    for i in range(0, x.shape[0]):
        y[i] = x[i] + e
        idx = (np.arange(i - n, i, 1, dtype=int) + 1)
        idx = np.delete(idx, idx < 0)
        e = predictor(y[idx])

    return y

##########################################
### Main Program  ########################
##########################################


document = Document()
if not Only_Tests:
    # generate raport
    document.add_heading('Report',0) # tworzenie nagłówków druga wartość to poziom nagłówka 
    document.add_paragraph("Autor: Artur Mizuła")

    document.add_section()
    document.add_heading('Wykresy testujące działanie algorytmów',1)

x=np.linspace(-1,1,1000)
y=0.9*np.sin(np.pi*x*4)

x_alaw_comp=A_law_compress(x)
x_alaw_comp=Kwant(x_alaw_comp,bit_test)
x_alaw_decomp=A_law_decompress(x_alaw_comp)

x_mulaw_comp=mu_law_compress(x)
x_mulaw_comp=Kwant(x_mulaw_comp,bit_test)
x_mulaw_decomp=mu_law_decompress(x_mulaw_comp)

y_alaw_decomp =A_law_decompress(Kwant(A_law_compress(y),bit_test))
y_mulaw_decomp =mu_law_decompress(Kwant(mu_law_compress(y),bit_test))

dpcm_c=DPCM_compress(y,bit_test)
dpcm_dec=DPCM_decompress(dpcm_c)
dpcm_c_p=DPCM_compress_pred(y,bit_test,n=DPCM_n,predictor=DPCM_predictor)
dpcm_dec_p=DPCM_decompress_pred(dpcm_c_p,n=DPCM_n,predictor=DPCM_predictor)

f1,axs=plt.subplots(1,2,num=1,figsize=(6,6)) 
f1.suptitle(f"Test kompresji law dla {bit_test} bitów")
axs[0].plot(x,x_alaw_comp,label="a_law")
axs[0].plot(x,x_mulaw_comp,label="mu_law")
axs[0].set_title("Sygnał po kompresji")
axs[0].legend()

axs[1].plot(x,x_alaw_decomp,label="a_law")
axs[1].plot(x,x_mulaw_decomp,label="mu_law")
axs[1].set_title("Sygnał po dekompresji")
axs[1].legend()

f2,axs=plt.subplots(5,1,num=2,figsize=(8,6)) 
f2.suptitle(f"Test dekompresji dla {bit_test} bitów")
axs[0].plot(x,y,label="Sygnał bazowy")
axs[0].set_title("Sygnał bazowy")

axs[1].plot(x,y_alaw_decomp,label="Sygnał po kompresji A-law")
axs[1].legend()

axs[2].plot(x,y_mulaw_decomp,label="Sygnał po kompresji mu-law")
axs[2].legend()

axs[3].plot(x,dpcm_dec,label="Sygnał po kompresji DPCM bez predykcji")
axs[3].legend()

axs[4].plot(x,dpcm_dec_p,label="Sygnał po kompresji DPCM z predykcją")
axs[4].legend()


if Only_Tests:
    plt.show()
else:
    memfile = BytesIO() 
    f1.savefig(memfile)
    document.add_picture(memfile, width=Inches(6)) # set document size
    memfile.close()
    f1.clf()
    memfile = BytesIO() 
    f2.savefig(memfile)
    document.add_picture(memfile, width=Inches(6)) # set document size
    memfile.close()
    f2.clf()  
    document.add_section()
    document.add_heading("Obserwacje na podstawie odsłuchanych plików ",1)
    document.add_paragraph("W metodach A-law i μ-law zmieniany jest sposób "
                           "reprezentacji amplitudy sygnału przed kwantyzacją. "
                           "Wartości sygnału w okolicach zera są odwzorowywane z większą dokładnością, natomiast "
                           "wartości bliskie maksymalnym amplitudom są odwzorowywane mniej dokładnie. Dzięki temu po kwantyzacji "
                           "więcej poziomów przypada na małe amplitudy, co zmniejsza błędy w tych obszarach. Na wykresach widać, że po dekompresji "
                           "sygnał przyjmuje postać schodkową, a największe odchylenia od oryginału pojawiają się dla dużych wartości sygnału.")
    document.add_paragraph("Metoda DPCM działa inaczej, ponieważ nie zapisuje bezpośrendio wartości sygnału, tylko różnice między kolejnymi"
                           "próbkami. W wersji z predykcją aktualna próbka jest zapisywana na podstawie kilku wcześniejszych wartości.Dzięki temu wartośći podawane kwantyzacji są mniejsze, co poprawia jakość sygnału po dekompresji.")
    document.add_heading("Zadanie 2.2",2)
    document.add_paragraph("Dla wszystkich analizowanych plików dźwiękowych przy kompresji do 8 bitów jakość była bardzo dobra i sygnał był poprawnie oraz wyraźnie odtwarzalny. Nie występowały zauważalne zniekształcenia ani utrata informacji.")
    document.add_heading("Zadanie 2.3", 2)
    document.add_paragraph("W tabeli przedstawiono ocenę jakości dźwięku po kompresji dla różnych liczby bitów.")



    oceny_23 = {
        "sing_high1.wav": [
            [8, "bardzo dobra", "bardzo dobra", "bardzo dobra", "bardzo dobra"],
            [7, "dobra", "dobra", "dobra", "bardzo dobra"],
            [6, "dobra", "dobra", "dobra", "dobra"],
            [5, "słaba", "średnia", "słaba", "słaba"],
            [4, "bardzo słaba", "bardzo slaba","trudno rozpoznać", "trudno rozpoznac"],
            [3, "trudno rozpoznać", "bardzo słaba", "trudno rozpoznac", "trudno rozpoznac"],
            [2, "trudno rozpoznać", "trudno rozpoznać", "trudno rozpoznac", "nierozpoznawalne"],
        ],
        "sing_medium1.wav": [
            [8, "bardzo dobra", "bardzo dobra", "bardzo dobra", "bardzo dobra"],
            [7, "bardzo dobra", "bardzo dobra", "dobra", "dobra"],
            [6, "dobra", "dobra", "dobra", "dobra"],
            [5, "średnia", "średnia", "srednia", "srednia"],
            [4, "średnia", "średnia", "średnia", "średnia"],
            [3, "słaba", "słaba", "słaba", "słaba"],
            [2, "bardzo słaba", "bardzo słaba", "bardzo słaba", "bardzo słaba"],
        ],
        "sing_low1.wav": [
            [8, "bardzo dobra", "bardzo dobra", "bardzo dobra", "bardzo dobra"],
            [7, "bardzo dobra", "bardzo dobra", "bardzo dobra", "bardzo dobra"],
            [6, "dobra", "dobra", "dobra", "dobra"],
            [5, "srednia", "srednia", "srednia", "slaba"],
            [4, "słaba", "słaba", "słaba", "słaba"],
            [3, "słaba", "słaba", "słaba", "słaba"],
            [2, "bardzo słaba", "bardzo słaba", "bardzo słaba", "bardzo słaba"],
        ],
    }
    for nazwa_pliku, dane_23 in oceny_23.items():
        document.add_paragraph(f"Ocena jakości dla pliku {nazwa_pliku}:")
        tabela_23 = document.add_table(rows=1, cols=5)
        tabela_23.style = "Table Grid"
        hdr = tabela_23.rows[0].cells
        hdr[0].text = "Liczba bitów"
        hdr[1].text = "A-law"
        hdr[2].text = "μ-law"
        hdr[3].text = "DPCM bez predykcji"
        hdr[4].text = "DPCM z predykcją"

        for row in dane_23:
            cells = tabela_23.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)

    document.add_section()
    document.add_heading("Podsumowanie i Wnioski",1)
    document.add_paragraph("Metody A-law i μ-law dają bardzo podobne rezultaty  "
                           "odsłuchowe. W przypadku DPCM zauważono, że "
                           "przy mniejszej liczbie bitów wersja bez predykcji często daje "
                           "lepsze rezultaty niż wersja z predykcją. Ogólnie DPCM przy niskiej "
                           "liczbie bitów wypada gorzej niż metody A-law i μ-law. Najlepiej "
                           "odtwarzany był sygnał o klasie medium, natomiast"
                           " najgorzej sygnał o klasie high, który był najbardziej"
                           " podatny na zniekształcenia.")
    document.save(OutputRaportFile) 
    # Audio files Generator
    for file in SingFiles:
        Signal, Fs = sf.read(os.path.join(AudioDir, file), dtype='float32')
        sfile = file.split(os.sep)[-1].split('.')
        for bit in [8, 7, 6, 5, 4, 3, 2]:
            y_alaw_decomp = A_law_decompress(Kwant(A_law_compress(Signal), bit))
            y_mulaw_decomp = mu_law_decompress(Kwant(mu_law_compress(Signal), bit))

            dpcm_c = DPCM_compress(Signal, bit)
            dpcm_dec = DPCM_decompress(dpcm_c)
            dpcm_c_p = DPCM_compress_pred(Signal, bit, n=DPCM_n, predictor=DPCM_predictor)
            dpcm_dec_p = DPCM_decompress_pred(dpcm_c_p, n=DPCM_n, predictor=DPCM_predictor)

            sf.write(os.path.join(OutputFolder,f"{sfile[0]}_A_LAW_{bit}b.wav"),data=y_alaw_decomp,samplerate=Fs)
            sf.write(os.path.join(OutputFolder,f"{sfile[0]}_mu_LAW_{bit}b.wav"),data=y_mulaw_decomp,samplerate=Fs)
            sf.write(os.path.join(OutputFolder,f"{sfile[0]}_DPCM_bp_{bit}b.wav"),data=dpcm_dec,samplerate=Fs)
            sf.write(os.path.join(OutputFolder,f"{sfile[0]}_DPCM_zp_{bit}b.wav"),data=dpcm_dec_p,samplerate=Fs)

            