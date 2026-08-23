import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import cv2
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO

##########################################
### Settings #############################
##########################################

Test=False
ColorFit_Test=False


GrayScale_bits = [1,2,4]


pallet8 = np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 0.0, 1.0,],
        [0.0, 1.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [1.0, 0.0, 0.0,],
        [1.0, 0.0, 1.0,],
        [1.0, 1.0, 0.0,],
        [1.0, 1.0, 1.0,],
])
pallet16 =  np.array([
        [0.0, 0.0, 0.0,], 
        [0.0, 1.0, 1.0,],
        [0.0, 0.0, 1.0,],
        [1.0, 0.0, 1.0,],
        [0.0, 0.5, 0.0,], 
        [0.5, 0.5, 0.5,],
        [0.0, 1.0, 0.0,],
        [0.5, 0.0, 0.0,],
        [0.0, 0.0, 0.5,],
        [0.5, 0.5, 0.0,],
        [0.5, 0.0, 0.5,],
        [1.0, 0.0, 0.0,],
        [0.75, 0.75, 0.75,],
        [0.0, 0.5, 0.5,],
        [1.0, 1.0, 1.0,], 
        [1.0, 1.0, 0.0,]
])

Color_pallets=[
    pallet8,pallet16
]

M2=np.array([[0,8,2,10],
             [12,4,14,6],
             [3,11,1,9],
             [12,7,13,5]])

OutputRaportFile = "report.docx"

##########################################
### Data Set #############################
##########################################

ImgDir = r'/Users/mizula/ZUT/SYSTEMY_MULTIMEDIALNE/PythonProject'# Address of folder with files (do nor delete `r``)


GsImages = ["IMG_GS/GS_0001.tif", "IMG_GS/GS_0002.png", "IMG_GS/GS_0003.png"]
ColorImages = ["IMG_SMALL/SMALL_0001.tif", "IMG_SMALL/SMALL_0002.png", "IMG_SMALL/SMALL_0003.png", "IMG_SMALL/SMALL_0007.jpg"]
##########################################
### Functions to  ########################
##########################################

def imgToUint8(img):
    if np.issubdtype(img.dtype, np.floating):
        img = (img * 255).astype(np.uint8)
    return img


def imgToFloat(img):
    if (np.issubdtype(img.dtype, np.integer)):
        img=img/255
    return img


def colorFit(pixel,Pallet):
        number =np.linalg.norm(np.abs(Pallet-pixel),axis=1)
        return Pallet[np.argmin(number)]



def kwant_colorFit(img,Pallet):
        out_img = img.copy()
        height,width=img.shape[:2]

        for k in range(height):
            for w in range(width):
                tmp=colorFit(img[k,w],Pallet)
                if len(tmp)==1:
                    out_img[k,w]=tmp[0]
                else:
                    out_img[k,w]=tmp[:]
        return out_img.astype(img.dtype)


def dith_randm(img):
        out_img = img.copy()
        height,width=img.shape[:2]
        r=np.random.rand(height,width)
        for k in range(height):
            for w in range(width):
                if img[k,w]>=r[k,w]:
                    out_img[k,w]=1
                else:
                    out_img[k,w]=0

        return out_img.astype(img.dtype)

def dith_ordered(img,Pallet,r=1,M=M2):
        out_img = img.copy()
        height, width = img.shape[:2]
        n=M.shape[0]//2
        Mpre=(M+1)/(2*n)**2-0.5

        for k in range(height):
            for w in range(width):
                value=img[k,w]+r*Mpre[k%(n*2),w%(n*2)]
                tmp=colorFit(np.array([value]),Pallet)
                if len(tmp) == 1:
                    out_img[k, w] = tmp[0]
                else:
                    out_img[k, w] = tmp[:]
        return out_img.astype(img.dtype)


def dith_FS(img, Pallet):
    out_img = img.copy()
    height, width = img.shape[:2]

    for x in range(height):
        for y in range(width):
            oldpixel = out_img[x, y].copy()

            if img.ndim == 2:
                tmp = colorFit(np.array([oldpixel]), Pallet)
                out_img[x, y] = tmp[0]
                quant_error = oldpixel - tmp[0]
            else:
                tmp = colorFit(oldpixel, Pallet)
                out_img[x, y] = tmp
                quant_error = oldpixel - tmp

            if y + 1 < width:
                out_img[x, y + 1] = out_img[x, y + 1] + quant_error * 7 / 16
            if x + 1 < height and y - 1 >= 0:
                out_img[x + 1, y - 1] = out_img[x + 1, y - 1] + quant_error * 3 / 16
            if x + 1 < height:
                out_img[x + 1, y] = out_img[x + 1, y] + quant_error * 5 / 16
            if x + 1 < height and y + 1 < width:
                out_img[x + 1, y + 1] = out_img[x + 1, y + 1] + quant_error * 1 / 16

    return out_img.astype(img.dtype)


##########################################
### Main Program  ########################
##########################################

def process_and_plot_GS(img,bit,filename,counter,figsize=(12,12)):
        if len(img.shape)>2:
                img=img[:,:,0]
        palett=np.linspace(0,1,2**bit).reshape(-1,1)
        qwant_img=kwant_colorFit(img,palett)

        order_img=dith_ordered(img,palett)
        FS_img=dith_FS(img,palett)
        if bit==1:
            rand_img=dith_randm(img)
            f,axs=plt.subplots(2,3,num=counter,figsize=figsize)
            f.suptitle(f"{filename} Dithering 1-bit")
            axs[0,0].imshow(img,cmap="gray")
            axs[0,0].set_title("Oryginał")
            axs[0,0].set_axis_off()

            axs[1,0].remove()

            axs[0,2].imshow(rand_img,cmap="gray")
            axs[0,2].set_title("Dithering\n Losowy")
            axs[0,2].set_axis_off()

            axt=[axs[0,1],axs[1,1],axs[1,2]]
        else:
            f,axs=plt.subplots(1,4,num=counter,figsize=figsize) 
            f.suptitle(f"{filename} Dithering {bit}-bitów")
            axs[0].imshow(img,cmap="gray")
            axs[0].set_title("Oryginał")
            axs[0].set_axis_off()  
            axt=[axs[1],axs[2],axs[3]] 
            rand_img=0

        axt[0].imshow(qwant_img,cmap="gray")
        axt[0].set_title("Kwantyzacja")
        axt[0].set_axis_off()

        
        axt[1].imshow(order_img,cmap="gray")
        axt[1].set_title("Dithering\n Zorganizowany")
        axt[1].set_axis_off()

        
        axt[2].imshow(FS_img,cmap="gray")
        axt[2].set_title("Dithering\n Floyda-Steinberga")
        axt[2].set_axis_off()

        if Test:
                if bit==1:
                    print(f"Test of uniqe values counts:\n"+
                        f"Unique values in Pallet: {np.unique(palett).size}\n"+
                        f"Dithering Random: {np.unique(rand_img).size}\n"+
                        f"Dithering Ordered: {np.unique(order_img).size}\n"+
                        f"Dithering Floyd-Steinberg: {np.unique(FS_img).size}\n")
                else:
                    print(f"Test of uniqe values counts:\n"+
                        f"Unique values in Pallet: {np.unique(palett).size}\n"+
                        f"Dithering Ordered: {np.unique(order_img).size}\n"+
                        f"Dithering Floyd-Steinberg: {np.unique(FS_img).size}\n")
        return f
                    

def process_and_plot_Color(img,palett,filename,counter,figsize=(9,13)):
        if img.shape[2]>3:
                img=img[:,:,:3]
        qwant_img=kwant_colorFit(img,palett)

        order_img=dith_ordered(img,palett)
        FS_img=dith_FS(img,palett)

        f,axs=plt.subplots(4,1,num=counter,figsize=figsize)
        f.suptitle(f"{filename} Dithering {len(palett)} kolorów")
        axs[0].imshow(img)
        axs[0].set_title("Oryginał")
        axs[0].set_axis_off()  

        axs[1].imshow(qwant_img)
        axs[1].set_title("Kwantyzacja")
        axs[1].set_axis_off()  

        axs[2].imshow(order_img)
        axs[2].set_title("Dithering\n Zorganizowany")
        axs[2].set_axis_off()  

        axs[3].imshow(FS_img)
        axs[3].set_title("Dithering\n Floyda-Steinberga")
        axs[3].set_axis_off() 

        return f




if Test:
        if ColorFit_Test:
            paleta = np.linspace(0,1,3).reshape(3,1)
            print(f"0.43 -> {colorFit(0.43,paleta)}") 
            print(f"0.66 -> {colorFit(0.66,paleta)}") 
            print(f"0.8 -> {colorFit(0.8,paleta)}") 

            print(f"[0.25,0.25,0.5] 8 kolorów -> {colorFit(np.array([0.25,0.25,0.5]),pallet8)}")
            print(f"[0.25,0.25,0.5] 16 kolorów-> {colorFit(np.array([0.25,0.25,0.5]),pallet16)}")

              
        file=GsImages[0]
        img=imgToFloat(plt.imread(os.path.join(ImgDir,file)))
        counter=1
        for bit in GrayScale_bits:
               process_and_plot_GS(img=img,bit=bit,filename=file,counter=counter)
               counter+=1

        file=ColorImages[0]
        img=imgToFloat(plt.imread(os.path.join(ImgDir,file)))
        for palett in Color_pallets:
               process_and_plot_Color(img=img,palett=palett,filename=file,counter=counter)
               counter+=1
        
        plt.show()
else:
    # generate raport
    document = Document()
    document.add_heading('Report',0) # tworzenie nagłówków druga wartość to poziom nagłówka 
    document.add_paragraph("Autor:  Artur Mizuła ")

    document.add_section()
    document.add_heading("Test ditheringu na obrazach w skali odcieni szarości",1)
    counter = 1
    for file in GsImages:
        img=imgToFloat(plt.imread(os.path.join(ImgDir,file)))
        for bit in GrayScale_bits:
                f = process_and_plot_GS(img=img,bit=bit,filename=file,counter=counter)
                counter+=1
                memfile = BytesIO()
                f.savefig(memfile)
                document.add_picture(memfile, width=Inches(6)) # set document size
                memfile.close()
                f.clf()
    document.add_section()
    document.add_heading("Test ditheringu na obrazach kolorowych",1)
    for file in ColorImages:
        img=imgToFloat(plt.imread(os.path.join(ImgDir,file)))
        for palett in Color_pallets:
            f = process_and_plot_Color(img=img,palett=palett,filename=file,counter=counter)
            counter += 1
            memfile = BytesIO()
            f.savefig(memfile)
            document.add_picture(memfile, width=Inches(6)) # set document size
            memfile.close()
            f.clf()
    document.add_section()
    document.add_heading("Podsumowanie i wnioski",1)
    document.add_paragraph(
        "W przypadku obrazów w skali szarości zauważalna jest duża różnica między metodami. "
        "Kwantyzacja powoduje znaczną utratę szczegółów oraz powstawanie ostrych przejść między tonami. "
        "Różne rodzaje ditheringu redukują te ostre przejścia. "
        "Dithering losowy wprowadza silny szum, przez co obraz staje się ziarnisty i mniej czytelny. "
        "Dithering zorganizowany poprawia rozkład odcieni, jednak wprowadza widoczny wzór oraz sprawia, że obraz jest bardziej wyblakły. "
        "Najlepsze rezultaty uzyskano przy użyciu ditheringu Floyda-Steinberga – obraz zachowuje więcej szczegółów i wygląda najbardziej naturalnie. \n"
    

        "W przypadku obrazów kolorowych efekty również zależą od użytej palety barw. "
        "Paleta 16 kolorów daje lepsze rezultaty niż paleta 8 kolorów, co jest szczególnie widoczne przy kwantyzacji. "
        "Porównując te same obrazy w RGB i skali szarości można zauważyć, że obrazy 4-bitowe w skali szarości zachowują więcej szczegółów "
        "niż obrazy zredukowane do 16 kolorów. "
        "Wynika to z faktu, że w przypadku obrazu kolorowego redukowane są jednocześnie trzy kanały, co prowadzi do większej utraty informacji. "
        "Podobnie jak w przypadku obrazów w skali szarości, również dla obrazów kolorowych najlepsze rezultaty uzyskano przy użyciu ditheringu Floyda-Steinberga, "
        "natomiast pozostałe metody dają wyraźnie gorsze efekty wizualne.\n"
      
        "Zaobserwowałem również, że przy większej liczbie bitów (np. 4 bity) w obrazach w skali szarości sama kwantyzacja często daje lepsze rezultaty wizualne niż dithering zorganizowany, "
        "ponieważ nie wprowadza dodatkowego wzoru i obraz pozostaje bardziej czytelny. "
        "Natomiast w przypadku obrazów kolorowych dithering zorganizowany znacząco poprawia odbiór obrazu w porównaniu do samej kwantyzacji, "
        "ponieważ pozwala lepiej oddać przejścia kolorów."
    )
    document.save(OutputRaportFile) 