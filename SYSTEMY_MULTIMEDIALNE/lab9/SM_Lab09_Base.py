import cv2
import numpy as np
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO
##############################################################################
######   Konfiguracja       ##################################################
##############################################################################

kat=r'.'                                 # katalog z plikami wideodob
plik_options = ["clip_1.mp4", "clip_3.mp4"]# nazwa pliku
ile=20                                # ile klatek odtworzyć? <0 - całość
key_frame_counter=4                     # co która klatka ma być kluczowa i nie podlegać kompresji
plot_frames=np.array([3,7,11])           # automatycznie wyrysuj wykresy
auto_pause_frames=np.array([])        # automatycznie za pauzuj dla klatki
subsampling="4:2:0"                    # parametry dla chroma subsampling
dzielnik=4
subsampling_options = ["4:4:4", "4:2:2", "4:4:0", "4:2:0", "4:1:1", "4:1:0"]
dzielnik_options = [1, 2, 4, 8]
key_frame_counter_options = [2, 4, 8, 12, 16,20]
wyswietlaj_kaltki=False                # czy program ma wyświetlać klatki
ROI_by_file = {
    "clip_1.mp4": [[120, 350, 850, 1250]],
    "clip_3.mp4": [[90, 270, 160, 480]],
}

##############################################################################
####     Kompresja i dekompresja    ##########################################
##############################################################################
def byteRun_encode(data):
    data=data.astype(int)
    x = np.array([len(data.shape)])
    x = np.concatenate([x, data.shape])
    flat=data.flatten()
    n=len(flat)

    buffer=np.zeros((n*2))
    out_idx=0
    i=0
    while i<n:
        start_i = i
        if i<n-1 and flat[i]==flat[i+1]:
            count=1
            for j in range(i+1,n):
                if(flat[j] == flat[i]):
                    count+=1

                else:
                    break
            full_count=count
            while count > 128:
                buffer[out_idx] = -127
                buffer[out_idx + 1] = flat[i]
                out_idx += 2
                count -= 128

            buffer[out_idx]=1-count
            buffer[out_idx+1]=flat[i]
            out_idx+=2
            i+=full_count
        else:
            values=[flat[i]]
            for j in range(i+1,n):
                if j<n-1 and flat[j]==flat[j+1]:
                    break

                values.append(flat[j])
            count=len(values)
            idx=0
            while count > 128:
                buffer[out_idx] = 127
                out_idx += 1
                for k in range(128):
                    buffer[out_idx] = values[idx]
                    out_idx += 1
                    idx += 1
                count -= 128
            buffer[out_idx] = count-1
            out_idx+=1
            for k in range(count):
                buffer[out_idx] = values[idx]
                out_idx += 1
                idx += 1
            i+=len(values)

    compressed_data = buffer[:out_idx]
    result= np.concatenate([x,compressed_data])
    return result.astype(int)

def byteRun_decode(data):
    OGshape = data[1:int(data[0] + 1)]
    encoded = data[int(data[0]) + 1:]
    decoded=[]
    i=0
    while i<len(encoded):
        count=encoded[i]
        if count<0:

            value=encoded[i+1]
            decoded.extend([value] * (1-count))
            i+=2
        else:
            decoded.extend(encoded[i+1:i+1+count+1])
            i += 1+ count+1
    decoded = np.array(decoded)
    return decoded.reshape(OGshape).astype(int)
class data:
    def __init__(self):
        # w pełni skompresowane dane
        self.Y=None 
        self.Cb=None
        self.Cr=None 
        # dane bez kompresji strumieniowej w celu przyspieszenia obliczeń
        self.semi_Y=None
        self.semi_Cb=None
        self.semi_Cr=None

def Chroma_subsampling(L,subsampling):
    if subsampling == "4:2:2":
        return L[:, ::2]
    elif subsampling == "4:4:0":
        return L[::2,:]
    elif subsampling == "4:2:0":
        return L[::2,::2]
    elif subsampling == "4:1:1":
        return L[:,::4]
    elif subsampling == "4:1:0":
        return L[::2,::4]
    else:
        return L


def Chroma_resampling(L,subsampling):
    if subsampling == "4:2:2":
        return np.repeat(L, 2, axis=1)
    elif subsampling == "4:4:0":
        return np.repeat(L, 2, axis=0)
    elif subsampling == "4:2:0":
        L=np.repeat(L, 2, axis=1)
        L=np.repeat(L, 2, axis=0)
        return L
    elif subsampling == "4:1:1":
        return np.repeat(L, 4, axis=1)
    elif subsampling == "4:1:0":
        L= np.repeat(L, 4, axis=1)
        L=np.repeat(L, 2, axis=0)
        return L
    else:
        return L
        
def frame_image_to_class(frame,subsampling):
    Frame_class = data()
    Frame_class.Y=frame[:,:,0].astype(int)
    Frame_class.Cb=Chroma_subsampling(frame[:,:,2].astype(int),subsampling)
    Frame_class.Cr=Chroma_subsampling(frame[:,:,1].astype(int),subsampling)
    return Frame_class


def frame_layers_to_image(Y,Cr,Cb,subsampling):  
    Cb=Chroma_resampling(Cb,subsampling)
    Cr=Chroma_resampling(Cr,subsampling)
    return np.dstack([Y,Cr,Cb]).clip(0,255).astype(np.uint8)

def compress_KeyFrame(Frame_class, use_byteRun=False):
    KeyFrame = data()
    KeyFrame.semi_Y = Frame_class.Y
    KeyFrame.semi_Cb = Frame_class.Cb
    KeyFrame.semi_Cr = Frame_class.Cr

    if use_byteRun:
        KeyFrame.Y = byteRun_encode(KeyFrame.semi_Y)
        KeyFrame.Cb = byteRun_encode(KeyFrame.semi_Cb)
        KeyFrame.Cr = byteRun_encode(KeyFrame.semi_Cr)
    else:
        KeyFrame.Y = KeyFrame.semi_Y
        KeyFrame.Cb = KeyFrame.semi_Cb
        KeyFrame.Cr = KeyFrame.semi_Cr

    return KeyFrame

def decompress_KeyFrame(KeyFrame, use_byteRun=False):
    if use_byteRun:
        Y = byteRun_decode(KeyFrame.Y)
        Cb = byteRun_decode(KeyFrame.Cb)
        Cr = byteRun_decode(KeyFrame.Cr)

        KeyFrame.semi_Y = Y
        KeyFrame.semi_Cb = Cb
        KeyFrame.semi_Cr = Cr
    else:
        Y = KeyFrame.semi_Y
        Cb = KeyFrame.semi_Cb
        Cr = KeyFrame.semi_Cr
    frame_image = frame_layers_to_image(Y, Cr, Cb, subsampling)
    return frame_image

def compress_not_KeyFrame(Frame_class, KeyFrame, dzielnik,use_byteRun=False):
    Compress_data = data()

    Compress_data.semi_Y = ((Frame_class.Y - KeyFrame.semi_Y) / dzielnik).astype(int)
    Compress_data.semi_Cb = ((Frame_class.Cb - KeyFrame.semi_Cb) / dzielnik).astype(int)
    Compress_data.semi_Cr = ((Frame_class.Cr - KeyFrame.semi_Cr) / dzielnik).astype(int)

    if(use_byteRun):
        Compress_data.Y = byteRun_encode(Compress_data.semi_Y)
        Compress_data.Cb = byteRun_encode(Compress_data.semi_Cb)
        Compress_data.Cr = byteRun_encode(Compress_data.semi_Cr)
    else:
        Compress_data.Y = Compress_data.semi_Y
        Compress_data.Cb = Compress_data.semi_Cb
        Compress_data.Cr = Compress_data.semi_Cr
    return Compress_data

def decompress_not_KeyFrame(Compress_data, KeyFrame, dzielnik, use_byteRun=False):

    if use_byteRun:
        diff_Y = byteRun_decode(Compress_data.Y)
        diff_Cb = byteRun_decode(Compress_data.Cb)
        diff_Cr = byteRun_decode(Compress_data.Cr)
    else:
        diff_Y = Compress_data.semi_Y
        diff_Cb = Compress_data.semi_Cb
        diff_Cr = Compress_data.semi_Cr

    Y = KeyFrame.semi_Y + diff_Y * dzielnik
    Cb = KeyFrame.semi_Cb + diff_Cb * dzielnik
    Cr = KeyFrame.semi_Cr + diff_Cr * dzielnik

    return frame_layers_to_image(Y, Cr, Cb, subsampling)

def plotDiffrence(ReferenceFrame,DecompressedFrame,FrameYr,DecompYr,ROI):
    # bardzo słaby i sztuczny przykład wykorzystania tej opcji
    # przerobić żeby porównanie było dokonywane w RGB nie YCrCb i/lub zastąpić innym porównaniem
    # ROI - Region of Insert współrzędne fragmentu który chcemy przybliżyć i ocenić w formacie [w1,w2,k1,k2]
    fig, axs = plt.subplots(4,3)
    fig.set_size_inches(12,6)
    diff = np.abs(ReferenceFrame[ROI[0]:ROI[1], ROI[2]:ROI[3]].astype(float) - DecompressedFrame[
        ROI[0]:ROI[1], ROI[2]:ROI[3]].astype(float))
    diff = np.mean(diff, axis=2)
    diffY = np.abs(FrameYr[ROI[0]:ROI[1], ROI[2]:ROI[3]][:,:,0].astype(float) - DecompYr[
        ROI[0]:ROI[1], ROI[2]:ROI[3]][:,:,0].astype(float))

    diffCb = np.abs(FrameYr[ROI[0]:ROI[1], ROI[2]:ROI[3]][:, :, 1].astype(float) - DecompYr[
        ROI[0]:ROI[1], ROI[2]:ROI[3]][:, :, 1].astype(float))
    diffCr = np.abs(FrameYr[ROI[0]:ROI[1], ROI[2]:ROI[3]][:, :, 2].astype(float) - DecompYr[
        ROI[0]:ROI[1], ROI[2]:ROI[3]][:, :, 2].astype(float))

    # print(np.min(diff), np.max(diff))
    axs[0,0].imshow(ReferenceFrame[ROI[0]:ROI[1],ROI[2]:ROI[3]])
    axs[0,1].imshow(DecompressedFrame[ROI[0]:ROI[1],ROI[2]:ROI[3]])
    axs[0,2].imshow(diff, cmap="gray", vmin=0, vmax=10)
    axs[0, 2].set_title("RGB")
    axs[1, 0].imshow(FrameYr[ROI[0]:ROI[1], ROI[2]:ROI[3]][:,:,0])

    axs[1, 1].imshow(DecompYr[ROI[0]:ROI[1], ROI[2]:ROI[3]][:,:,0])

    axs[1,2].imshow(diffY, cmap="gray", vmin=0, vmax=10)
    axs[1, 2].set_title("Y")
    axs[2, 0].imshow(FrameYr[ROI[0]:ROI[1], ROI[2]:ROI[3]][:, :, 1])
    axs[2, 1].imshow(DecompYr[ROI[0]:ROI[1], ROI[2]:ROI[3]][:, :, 1])
    axs[2, 2].imshow(diffCb, cmap="gray", vmin=0, vmax=10)
    axs[2, 2].set_title("Cr")
    axs[3, 0].imshow(FrameYr[ROI[0]:ROI[1], ROI[2]:ROI[3]][:, :, 2])
    axs[3, 1].imshow(DecompYr[ROI[0]:ROI[1], ROI[2]:ROI[3]][:, :, 2])
    axs[3, 2].imshow(diffCr, cmap="gray", vmin=0, vmax=10)
    axs[3, 2].set_title("Cb")

    return fig


##############################################################################
####     Głowna pętla programu      ##########################################
##############################################################################
document = Document()
document.add_heading("Lab09",0)
image_counter = 0
for subsampling in subsampling_options:
    for dzielnik in dzielnik_options:
        for plik in plik_options:
            cap = cv2.VideoCapture(os.path.join(kat,plik))

            if ile<0:
                ile=int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

            cv2.namedWindow('Normal Frame')
            cv2.namedWindow('Decompressed Frame')

            compression_information=np.zeros((3,ile))

            for i in range(ile):
                ret, frame = cap.read()
                frame_RGB = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                if wyswietlaj_kaltki:
                    cv2.imshow('Normal Frame',frame)

                frame=cv2.cvtColor(frame,cv2.COLOR_BGR2YCrCb)
                Frame_class = frame_image_to_class(frame,subsampling)
                if (i % key_frame_counter)==0: # pobieranie klatek kluczowych
                    KeyFrame = compress_KeyFrame(Frame_class)
                    cY=KeyFrame.Y
                    cCb=KeyFrame.Cb
                    cCr=KeyFrame.Cr
                    Decompresed_Frame = decompress_KeyFrame(KeyFrame)
                    decomp_RGB = cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2RGB)
                else: # kompresja
                    Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame,dzielnik)
                    cY=Compress_data.Y
                    cCb=Compress_data.Cb
                    cCr=Compress_data.Cr
                    Decompresed_Frame = decompress_not_KeyFrame(Compress_data,  KeyFrame,dzielnik)
                    decomp_RGB = cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2RGB)
                compression_information[0,i]= (frame[:,:,0].size - cY.size)/frame[:,:,0].size
                compression_information[1,i]= (frame[:,:,0].size - cCb.size)/frame[:,:,0].size
                compression_information[2,i]= (frame[:,:,0].size - cCr.size)/frame[:,:,0].size
                if wyswietlaj_kaltki:
                    cv2.imshow('Decompressed Frame',cv2.cvtColor(Decompresed_Frame,cv2.COLOR_YCrCb2BGR))

                if np.any(plot_frames==i): # rysuj wykresy
                    for r in ROI_by_file[plik]:
                        fig = plotDiffrence(frame_RGB,decomp_RGB,frame,Decompresed_Frame,r)
                        document.add_paragraph(f"Plik={plik}, Subsampling={subsampling}, dzielnik={dzielnik}, frame={i}")
                        memfile = BytesIO()
                        fig.savefig(memfile, format="png", bbox_inches="tight")
                        memfile.seek(0)
                        document.add_picture(memfile, width=Inches(4.5))

                        image_counter += 1
                        if image_counter % 6 == 0:
                            document.add_page_break()
                        memfile.close()
                        plt.close(fig)
                if np.any(auto_pause_frames==i):
                    cv2.waitKey(-1) #wait until any key is pressed

                k = cv2.waitKey(1) #& 0xff

                if k==ord('q'):
                    break
                elif k == ord('p'):
                    cv2.waitKey(-1) #wait until any key is pressed

            # plt.figure()
            # plt.plot(np.arange(0,ile),compression_information[0,:]*100)
            # plt.plot(np.arange(0,ile),compression_information[1,:]*100)
            # plt.plot(np.arange(0,ile),compression_information[2,:]*100)
            # plt.title("File:{}, subsampling={}, divider={}, KeyFrame={} ".format(plik,subsampling,dzielnik,key_frame_counter))
            # plt.show()
            cap.release()
document.add_paragraph(f"Badanie skuteczności kompresji z użyciem RLE lub ByteRun")
subsampling = "4:2:0"
dzielnik = 4
use_byteRun = True
for key_frame in key_frame_counter_options:
    for plik in plik_options:
        cap = cv2.VideoCapture(os.path.join(kat, plik))

        if ile < 0:
            ile = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

        cv2.namedWindow('Normal Frame')
        cv2.namedWindow('Decompressed Frame')

        compression_information = np.zeros((3, ile))
        for i in range(ile):
            ret, frame = cap.read()
            frame_RGB = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            frame = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
            Frame_class = frame_image_to_class(frame, subsampling)

            if (i % key_frame) == 0:  # pobieranie klatek kluczowych
                KeyFrame = compress_KeyFrame(Frame_class,True)
                cY = KeyFrame.Y
                cCb = KeyFrame.Cb
                cCr = KeyFrame.Cr
                Decompresed_Frame = decompress_KeyFrame(KeyFrame,True)
                decomp_RGB = cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2RGB)
            else:  # kompresja
                Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame, dzielnik,True)
                cY = Compress_data.Y
                cCb = Compress_data.Cb
                cCr = Compress_data.Cr
                Decompresed_Frame = decompress_not_KeyFrame(Compress_data, KeyFrame, dzielnik,True)
                decomp_RGB = cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2RGB)
            compression_information[0, i] = (frame[:, :, 0].size - cY.size) / frame[:, :, 0].size
            compression_information[1, i] = (frame[:, :, 0].size - cCb.size) / frame[:, :, 0].size
            compression_information[2, i] = (frame[:, :, 0].size - cCr.size) / frame[:, :, 0].size
            if np.any(plot_frames == i):  # rysuj wykresy
                for r in ROI_by_file[plik]:
                    fig = plotDiffrence(frame_RGB, decomp_RGB,frame,Decompresed_Frame,r)
                    document.add_paragraph(f"Plik={plik}, Byterun=True, Subsampling={subsampling}, dzielnik={dzielnik}, key_frame={key_frame}, frame={i}")
                    memfile = BytesIO()
                    fig.savefig(memfile, format="png", bbox_inches="tight")
                    memfile.seek(0)
                    document.add_picture(memfile, width=Inches(4.5))
                    image_counter += 1
                    memfile.close()
                    plt.close(fig)
            if np.any(auto_pause_frames == i):
                cv2.waitKey(-1)  # wait until any key is pressed

            k = cv2.waitKey(1)  # & 0xff

            if k == ord('q'):
                break
            elif k == ord('p'):
                cv2.waitKey(-1)  # wait until any key is pressed

        fig = plt.figure(figsize=(6, 4))
        plt.plot(np.arange(0, ile), compression_information[0, :] * 100, label="Y")
        plt.plot(np.arange(0, ile), compression_information[1, :] * 100, label="Cb")
        plt.plot(np.arange(0, ile), compression_information[2, :] * 100, label="Cr")
        plt.title(f"File={plik}, ByteRun=True, subsampling={subsampling}, divider={dzielnik}, key_frame={key_frame}")

        memfile = BytesIO()
        fig.savefig(memfile, format="png", bbox_inches="tight")
        memfile.seek(0)

        document.add_paragraph(f"Wykres ByteRun: plik={plik}, key_frame={key_frame}")
        document.add_picture(memfile, width=Inches(4.8))

        document.add_page_break()

        memfile.close()
        plt.close(fig)
        cap.release()
document.add_heading("Podsumowanie i wnioski", 1)
document.add_paragraph(
    "Jako kompromis między jakością a kompresją wybrałem ustawienie 4:2:0 oraz dzielnik 4. "
    "Na mapach różnic widoczne były szare obszary, co oznaczało występowanie błędów rekonstrukcji, jednak na obrazie różnice nie były bardzo widoczne bardzo mocno. "
    "Ten wariant uznałem za dobry kompromis między utratą jakości a zmniejszeniem ilości danych."
    "Największy wpływ na pogorszenie jakości miał dzielnik. Wraz ze wzrostem wartości dzielnika różnice między klatką oryginalną a zdekompresowaną były coraz większe. "
    "Najgorsze rezultaty jakościowe uzyskano dla wariantu 4:1:0 oraz dzielnika 8. "
    "Największy problem przy rekonstrukcji był widoczny na zielonych ramkach wykrywających auta. Ramki te czasami dublowały się lub nie odtwarzały się poprawnie, ponieważ są to ostre, kontrastowe elementy"

)
document.add_paragraph(
    "W drugiej części ćwiczenia zastosowałem kompresję ByteRun. Na wykresach było widać, że dla kanału Y zysk pamięci był dużo mniejszy niż dla kanałów Cb i Cr. "
    "Dla kanałów Cb i Cr wyniki były dużo lepsze około 90% zysku pamięci. "
    "Na obrazach różnic pojawiały się szare fragmenty, czyli błędy rekonstrukcji, ale przy normalnym oglądaniu obrazu bez dużego przybliżania różnice nie były bardzo widoczne. "
    "Czym większy odstęp między klatkami kluczowymi na wykresach tym mniej częstych spadków, ponieważ pełne klatki kluczowe występowały rzadziej. "
    "Według mojej opinii najbardziej opłacało się kompresować kanał Y dla klatek niekluczowych. Widać wyraźnie na wykresie, że zysk pamięci spadał przy klatkach kluczowych. "
    "Kanały Cb i Cr kompresowały się dobrze zarówno dla klatek kluczowych, jak i niekluczowych.Można używać kompresji przy obu."
)


document.save("raport.docx")