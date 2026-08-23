import cv2
import numpy as np
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO
from skimage.metrics import structural_similarity as ssim
img1=cv2.imread("img1.jpg")
img2=cv2.imread("img2.jpg")
img3=cv2.imread("img3.jpg")
print(img1.shape)

print(img2.shape)
binear=cv2.imread("binear.png")
def water_mark(img,mask,alpha=0.25):
    assert (img.shape[0]==mask.shape[0]) and (img.shape[1]==mask.shape[1]), "Wrong size"
    if len(img.shape)<3:
        flag=True
        t_img=cv2.cvtColor(img,cv2.COLOR_GRAY2RGBA)
    else:
        flag=False
        t_img=cv2.cvtColor(img,cv2.COLOR_RGB2RGBA)
    if (mask.dtype==bool):
        t_mask=cv2.cvtColor((mask*255).astype(np.uint8),cv2.COLOR_GRAY2RGBA)
    elif (mask.dtype==np.uint8):
        if len(mask.shape)<3:
            t_mask=cv2.cvtColor((mask).astype(np.uint8),cv2.COLOR_GRAY2RGBA)
        else:
            t_mask=cv2.cvtColor((mask).astype(np.uint8),cv2.COLOR_RGB2RGBA)
    else:
        if len(mask.shape)<3:
            t_mask=cv2.cvtColor((mask*255).astype(np.uint8),cv2.COLOR_GRAY2RGBA)
        else:
            t_mask=cv2.cvtColor((mask*255).astype(np.uint8),cv2.COLOR_RGB2RGBA)
    t_out=cv2.addWeighted(t_img,1,t_mask,alpha,0)
    if flag:
        out=cv2.cvtColor(t_out,cv2.COLOR_RGBA2GRAY)
    else:
        out=cv2.cvtColor(t_out,cv2.COLOR_RGBA2RGB)
    return out

def put_data(img,data,binary_mask=np.uint8(1)):
    assert img.dtype==np.uint8 , "img wrong data type"
    assert binary_mask.dtype==np.uint8, "binary_mask wrong data type"
    un_binary_mask=np.unpackbits(binary_mask)
    if data.dtype!=bool:
        unpacked_data=np.unpackbits(data)
    else:
        unpacked_data=data
    dataspace=img.shape[0]*img.shape[1]*np.sum(un_binary_mask)
    assert (dataspace>=unpacked_data.size) , "too much data"
    if dataspace==unpacked_data.size:
        prepered_data=unpacked_data.reshape(img.shape[0],img.shape[1],int(np.sum(un_binary_mask))).astype(np.uint8)
    else:
        prepered_data=np.resize(unpacked_data,(img.shape[0],img.shape[1],int(np.sum(un_binary_mask)))).astype(np.uint8)
    mask=np.full((img.shape[0],img.shape[1]),binary_mask)
    img=np.bitwise_and(img,np.invert(mask))
    bv=0
    for i,b in enumerate(un_binary_mask[::-1]):
        if b: 
            temp=prepered_data[:,:,bv]
            temp=np.left_shift(temp,i)
            img=np.bitwise_or(img,temp)
            bv+=1
    return img

def pop_data(img,binary_mask=np.uint8(1),out_shape=None):
    un_binary_mask=np.unpackbits(binary_mask)
    data=np.zeros((img.shape[0],img.shape[1],np.sum(un_binary_mask))).astype(np.uint8)
    bv=0
    for i,b in enumerate(un_binary_mask[::-1]):
        if b:
            mask=np.full((img.shape[0],img.shape[1]),2**i)
            temp=np.bitwise_and(img,mask)
            data[:,:,bv]=temp[:,:].astype(np.uint8)
            bv+=1
    if out_shape!=None:
        tmp=np.packbits(data.flatten())
        tmp=tmp[:np.prod(out_shape)]
        data=tmp.reshape(out_shape)
    return data
def MSE(image1, image2):
    image1 = image1.astype(np.float32)
    image2 = image2.astype(np.float32)
    return np.mean((image1 - image2) ** 2)
def PSNR(image1, image2):
    image1 = image1.astype(np.float32)
    image2 = image2.astype(np.float32)
    return 10*np.log10(255**2/(MSE(image1, image2)+0.01))
def SSIM(image1, image2):
    image1 = image1.astype(np.float32)
    image2 = image2.astype(np.float32)
    if len(image1.shape) == 2:
        return ssim(image1, image2, data_range=255)
    else:
        return ssim(image1, image2, channel_axis=2, data_range=255)
#Zad1
document = Document()
document.add_heading("Lab11",0)
document.add_heading("Zadanie1 - ukrywanie tekstu w kanale B", 1)
text = "Systemy multimedalne Artur Mizula"

A = np.array([ord(c) for c in list(text)])
data = np.frombuffer(A.tobytes(), dtype=np.uint8)

blue = img1[:, :, 0].copy()
blue_hidden = put_data(blue, data, binary_mask=np.uint8(1))

img_hidden = img1.copy()
img_hidden[:, :, 0] = blue_hidden

recovered = pop_data(img_hidden[:, :, 0],binary_mask=np.uint8(1),out_shape=data.shape)

A_rec = np.frombuffer(recovered.tobytes(), dtype=A.dtype)
recovered_text = ''.join(chr(x) for x in A_rec)
psnr_value = PSNR(blue, blue_hidden)
ssim_value = SSIM(blue, blue_hidden)
is_correct = text == recovered_text

fig, ax = plt.subplots(1, 2, figsize=(10, 5))

ax[0].imshow(cv2.cvtColor(img1, cv2.COLOR_BGR2RGB))
ax[0].set_title("Obraz oryginalny")
ax[0].axis("off")

ax[1].imshow(cv2.cvtColor(img_hidden, cv2.COLOR_BGR2RGB))
ax[1].set_title("Obraz z ukrytym tekstem")
ax[1].axis("off")

plt.tight_layout()
plt.savefig("zad1_porownanie.png")
plt.close()
document.add_picture("zad1_porownanie.png", width=Inches(5))
document.add_paragraph(f"Tekst wejściowy: {text}")
document.add_paragraph(f"Tekst odzyskany: {recovered_text}")
document.add_paragraph(f"PSNR: {psnr_value:.4f}")
document.add_paragraph(f"SSIM: {ssim_value:.4f}")
#Zad2
img1_rgb = cv2.cvtColor(img_hidden, cv2.COLOR_BGR2RGB)
binear_gray =cv2.imread('binear.png',cv2.IMREAD_GRAYSCALE)
fig, ax = plt.subplots(1, 4, figsize=(16, 5))
ax[0].imshow(img1_rgb)
ax[0].set_title("Oryginał")
ax[0].axis("off")

results = []
for i, alpha in enumerate([0.10, 0.25, 0.50]):
    wm = water_mark(img1_rgb, binear_gray, alpha)
    psnr_value = PSNR(wm, img1_rgb)
    ssim_value = SSIM(wm, img1_rgb)
    results.append((alpha, psnr_value, ssim_value))
    ax[i + 1].imshow(wm)
    ax[i + 1].set_title(f"alpha={alpha}")
    ax[i + 1].axis("off")
plt.tight_layout()
plt.savefig("zad2_watermark.png")
plt.close()
document.add_heading("Zadanie 2 - znak wodny", 1)
document.add_picture("zad2_watermark.png", width=Inches(6))
document.add_paragraph("Rysunek 2. Wpływ parametru alpha na widoczność znaku wodnego.")
table = document.add_table(rows=1, cols=3)
table.style = "Table Grid"

hdr = table.rows[0].cells
hdr[0].text = "Alpha"
hdr[1].text = "PSNR"
hdr[2].text = "SSIM"

for alpha, psnr_value, ssim_value in results:
    row = table.add_row().cells
    row[0].text = f"{alpha:.2f}"
    row[1].text = f"{psnr_value:.4f}"
    row[2].text = f"{ssim_value:.4f}"
#Zad3
document.add_heading("Zadanie 3 - ukrywanie obrazu w obrazie", 1)
masks_B = [7, 15, 31, 63, 127]
masks_G = [3, 7, 15, 31, 63]
masks_R = [3, 7, 15, 31]
hidden_img = img2.copy()
carrier = img1.copy()
current_masks=[7,3,3]
results_zad3 = []
for i in range(14):
    result = carrier.copy()
    channel=i%3
    step=i//3
    if channel==0:
        current_masks[0] = masks_B[step]
    elif channel == 1:
        current_masks[1] = masks_G[step]
    else:
        current_masks[2] = masks_R[step]
    result[:, :, 0] = put_data(carrier[:, :, 0], hidden_img[:, :, 0], np.uint8(current_masks[0]))  # B - 3 bity
    result[:, :, 1] = put_data(carrier[:, :, 1], hidden_img[:, :, 1], np.uint8(current_masks[1]))  # G - 2 bity
    result[:, :, 2] = put_data(carrier[:, :, 2], hidden_img[:, :, 2], np.uint8(current_masks[2]))  # R - 2 bity
    psnr_value = PSNR(carrier, result)
    ssim_value = SSIM(carrier, result)
    recovered_img = np.zeros_like(hidden_img)

    recovered_img[:, :, 0] = pop_data(result[:, :, 0], np.uint8(current_masks[0]), out_shape=hidden_img[:, :, 0].shape)
    recovered_img[:, :, 1] = pop_data(result[:, :, 1], np.uint8(current_masks[1]), out_shape=hidden_img[:, :, 1].shape)
    recovered_img[:, :, 2] = pop_data(result[:, :, 2], np.uint8(current_masks[2]), out_shape=hidden_img[:, :, 2].shape)
    results_zad3.append({
        "result": result.copy(),
        "recovered": recovered_img.copy(),
        "masks": current_masks.copy(),
        "psnr": psnr_value,
        "ssim": ssim_value

    })
selected = [
    results_zad3[0],
    results_zad3[len(results_zad3)//2],
    results_zad3[-1]
]
fig, ax = plt.subplots(2, 3, figsize=(15, 10))

for idx, item in enumerate(selected):

    b, g, r = item["masks"]

    ax[0, idx].imshow(cv2.cvtColor(item["result"], cv2.COLOR_BGR2RGB))
    ax[0, idx].set_title(
        f"B={b} G={g} R={r}\n"
        f"PSNR={item['psnr']:.2f}\n"
        f"SSIM={item['ssim']:.4f}"
    )
    ax[0, idx].axis("off")

    ax[1, idx].imshow(cv2.cvtColor(item["recovered"], cv2.COLOR_BGR2RGB))
    ax[1, idx].set_title("Odzyskany obraz")
    ax[1, idx].axis("off")

plt.tight_layout()
plt.savefig("zad3_porownanie.png")
plt.close()
document.add_picture("zad3_porownanie.png", width=Inches(6))
document.add_paragraph(
    "Rysunek 3. Porównanie obrazu z ukrytym obrazem oraz obrazu odzyskanego dla różnych masek bitowych."
)


document.add_paragraph("Przypadek 1:")
b, g, r = selected[0]["masks"]
document.add_paragraph(f"B={b}, G={g}, R={r} | "f"PSNR={selected[0]['psnr']:.4f} | "f"SSIM={selected[0]['ssim']:.4f}")

document.add_paragraph("Przypadek 2:")
b, g, r = selected[1]["masks"]
document.add_paragraph(f"B={b}, G={g}, R={r} | "f"PSNR={selected[1]['psnr']:.4f} | "f"SSIM={selected[1]['ssim']:.4f}")

document.add_paragraph("Przypadek 3:")
b, g, r = selected[2]["masks"]
document.add_paragraph(f"B={b}, G={g}, R={r} | "f"PSNR={selected[2]['psnr']:.4f} | "f"SSIM={selected[2]['ssim']:.4f}")
document.add_heading("Obserwacje i wnioski", 1)
document.add_paragraph("Ukrywanie tekstu w obrazie")
document.add_paragraph("Ukrycie tekstu w najmłodszym bicie kanału niebieskiego praktycznie nie wpływa na jakość obrazu. "
                       "Wartości PSNR i SSIM były bardzo wysokie. Różnice były niezauważalne.")
document.add_paragraph("Znak wodny")

document.add_paragraph("Dla wartości α = 0.10 znak wodny był słabo widoczny, natomiast dla α = 0.50 znak wodny zakrywał "
                       "oryginalny obraz. Najlepszym parametrem była wartość α = 0.25.")
document.add_paragraph("Ukrywanie obrazu w obrazie")


document.add_paragraph("W zadaniu stopniowo zwiększałem liczbę używanych bitów w kolejnych kanałach RGB. "
                       "W każdej iteracji zwiększałem liczbę wykorzystanych bitów w następnym kanale. "
                       "Pierwsze wyraźnie zauważalne przebarwienia oraz utrata jakości pojawiły się dla konfiguracji "
                       "B = 127, G = 63, R = 31 (odpowiednio 7 bitów, 6 bitów i 5 bitów). W tym przypadku wartości metryk osiągnęły niskie wartości.")
document.save("raport_lab13.docx")