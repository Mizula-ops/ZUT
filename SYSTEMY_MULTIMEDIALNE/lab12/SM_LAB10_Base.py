import cv2
import numpy as np
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO
from skimage.metrics import structural_similarity as ssim
photo = ["img1.jpg", "img2.jpg", "img3.jpg","img4.jpg"]

def jpegCompression(image,quality=70):
    encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), quality]
    result, encimg = cv2.imencode('.jpg', image, encode_param)
    decimg = cv2.imdecode(encimg, 1)
    return decimg

def blurImage(img, kernal_size=(5,5), method='gaussian', d=5, sigmaColor=75, sigmaSpace=75):
    if method=='blur':
        img = cv2.blur(img, kernal_size)
    elif method=='gaussian':
        img = cv2.GaussianBlur(img,kernal_size,sigmaX=5,sigmaY=5)
    elif method=='medianBlur':
        img= cv2.medianBlur(img, 5)
    elif method=='bilateral':
        img = cv2.bilateralFilter(img, d, sigmaColor, sigmaSpace)
    return img
def noiseImage(img,sigma=25,alpha=0.5):
    gauss = np.random.normal(0, sigma, img.shape)
    noisy1 = (img + alpha * gauss).clip(0, 255).astype(np.uint8)
    return noisy1
def MSE(image1, image2):
    image1 = image1.astype(np.float32)
    image2 = image2.astype(np.float32)
    m, n, c = image1.shape
    return np.sum((image1 - image2) ** 2) / (m * n* c)
def NMSE(image1, image2):
    image1 = image1.astype(np.float32)
    image2 = image2.astype(np.float32)
    c,n,m = image2.shape
    return MSE(image1, image2)/(np.mean(image2**2))
    # return np.sum((image1 - image2) ** 2) / np.sum(image1 ** 2)
def PSNR(image1, image2):
    image1 = image1.astype(np.float32)
    image2 = image2.astype(np.float32)
    return 10*np.log10(255**2/(MSE(image1, image2)+0.01))
def IF(image1, image2):
    image1 = image1.astype(np.float32)
    image2 = image2.astype(np.float32)

    return 1 - (np.sum((image2 - image1) ** 2)/ np.sum(image1 * image2))
def SSIM(image1, image2):
    image1 = image1.astype(np.float32)
    image2 = image2.astype(np.float32)
    return ssim(image1, image2,channel_axis=2,data_range=255)
img1 = cv2.imread(photo[0])
img2 = cv2.imread(photo[1])
img3 = cv2.imread(photo[2])
img4 = cv2.imread(photo[3])
images1=[img1]
images2=[img2]
images3=[img3]
images4=[img4]
images=[images1,images2,images3,images4]

jpeg_quality = [90,80,70,60,50,40,30,20]
blur_kernel = [3,5,7,9,11,13,15,17]
bilateral_d = [3,5,7,9,11,13,15,17]
noise_sigma = [5,10,15,20,25,30,35,40]

x = range(8)
names = [
    "JPEG Compression",
    "Gaussian Blur",
    "Bilateral Blur",
    "Gaussian Noise"
]

document = Document()
document.add_heading("Lab10",0)
for j in range(8):

    temp = img1.copy()
    for _ in range(5):
        temp = jpegCompression(temp, jpeg_quality[j])
    images1.append(temp)

    temp = img2.copy()
    k = blur_kernel[j]
    for _ in range(5):
        temp = blurImage(temp, (k, k), 'gaussian')
    images2.append(temp)

    temp = img3.copy()
    d = bilateral_d[j]
    for _ in range(5):
        temp = blurImage(temp, method='bilateral', d=d)
    images3.append(temp)

    temp = img4.copy()
    for _ in range(5):
        temp = noiseImage(temp, sigma=noise_sigma[j])
    images4.append(temp)
document.add_heading("Zdjecia", 1)
for idx, ims in enumerate(images):
    document.add_paragraph(f"Badany obraz: {photo[idx]}")
    document.add_paragraph(f"Metoda degradacji: {names[idx]}")
    MSE_score=[]
    NMSE_score=[]
    SSIM_score=[]
    PSNR_score=[]
    IF_score=[]
    fig, axs = plt.subplots(2, 4)
    fig.set_size_inches(16, 5)
    axs[0,0].imshow(ims[1])
    axs[0,0].set_title(f"{names[idx]} Oryginal")
    xd=0
    original_image = ims[0]

    for i, im in enumerate(ims[1:]):
        MSE_score.append(MSE(original_image, im))
        NMSE_score.append(NMSE(original_image, im))
        SSIM_score.append(SSIM(original_image, im))
        PSNR_score.append(PSNR(original_image, im))
        IF_score.append(IF(original_image, im))
        row = i // 4
        col = i % 4
        axs[row, col].imshow(cv2.cvtColor(im, cv2.COLOR_BGR2RGB))
        axs[row, col].set_title(f"{names[idx]} próba {i + 1}")
        axs[row, col].axis("off")

    memfile = BytesIO()
    fig.savefig(memfile, format="png", bbox_inches="tight")
    memfile.seek(0)
    document.add_picture(memfile, width=Inches(4.5))
    metrics = [
        ("MSE", MSE_score),
        ("NMSE", NMSE_score),
        ("SSIM", SSIM_score),
        ("PSNR", PSNR_score),
        ("IF", IF_score)
    ]
    fig, axs = plt.subplots(1, 5)
    fig.set_size_inches(16, 5)
    if idx == 0:
        x_values = jpeg_quality
    elif idx == 1:
        x_values = blur_kernel
    elif idx == 2:
        x_values = bilateral_d
    else:
        x_values = noise_sigma
    for i, (metric_name, values) in enumerate(metrics):
        axs[i].plot(x_values, values, marker='o')
        axs[i].set_title(f"{names[idx]} - {metric_name}")
        axs[i].set_xlabel("Wartość parametru")
        axs[i].set_ylabel(metric_name)
        axs[i].set_xticks(x_values)
        axs[i].grid()

    memfile = BytesIO()
    fig.savefig(memfile, format="png", bbox_inches="tight")
    memfile.seek(0)
    document.add_picture(memfile, width=Inches(4.5))
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "MSE"
    hdr[2].text = "NMSE"
    hdr[3].text = "SSIM"
    hdr[4].text = "PSNR"
    hdr[5].text = "IF"

    for i in range(8):
        row = table.add_row().cells

        if idx == 0:
            row[0].text = str(jpeg_quality[i])
        elif idx == 1:
            row[0].text = str(blur_kernel[i])
        elif idx == 2:
            row[0].text = str(bilateral_d[i])
        else:
            row[0].text = str(noise_sigma[i])
        row[1].text = f"{MSE_score[i]:.4f}"
        row[2].text = f"{NMSE_score[i]:.8f}"
        row[3].text = f"{SSIM_score[i]:.4f}"
        row[4].text = f"{PSNR_score[i]:.4f}"
        row[5].text = f"{IF_score[i]:.4f}"
    document.add_page_break()
document.add_heading("Podsumowanie i wnioski", 1)
document.add_paragraph("KOMPRESJA JPEG")
document.add_paragraph("Wraz ze zmiejszaniem jakości rosły wartości MSe i NMSE, co oznacza większy błąd, różnice względem oryginału."
                       "Jednocześnie wartości SSIM, PSNR i IF stopniowo malały. Czym mniejsze quaility tym większe pogorszenie")
document.add_paragraph("Rozmycie Gaussowskie")
document.add_paragraph("Im większy był parametr kernel, tym silniejsze było rozmycie obrazu.")
document.add_paragraph("Rozmycie biletaralne")
document.add_paragraph("tutaj zmieniano parametr d. Wraz ze wzrostem tego parametru obraz był coraz mocniej wygładzany. Ta metoda w prorównaniu z rozmyciem "
                       "Gaussowskim lepiej zachowuje krawędzie pomimo pogorszenia jakości."
                       "Przy rozmyciach występowały najbardziej regularne zmian wraz ze zmianą parametru.")
document.add_paragraph("Szum ")
document.add_paragraph("Największe pogorszenie jakości uzyskano dla szumu.MSE wzrosło do ponad 1350, a SSIm spadło do około 0.12 co oznacza bardzo duża utratę podobieństwa do oryginału. ")

document.save("raport1.docx")