import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import cv2
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO
import scipy.fftpack
##########################################
### Settings #############################
##########################################

Test=False

OutputRaportFile = "yes1.docx"

Chroma_options=["4:4:4","4:2:2"]
Quant_options=[True,False]

QY= np.array([
        [16, 11, 10, 16, 24,  40,  51,  61],
        [12, 12, 14, 19, 26,  58,  60,  55],
        [14, 13, 16, 24, 40,  57,  69,  56],
        [14, 17, 22, 29, 51,  87,  80,  62],
        [18, 22, 37, 56, 68,  109, 103, 77],
        [24, 36, 55, 64, 81,  104, 113, 92],
        [49, 64, 78, 87, 103, 121, 120, 101],
        [72, 92, 95, 98, 112, 100, 103, 99],
        ])

QC= np.array([
        [17, 18, 24, 47, 99, 99, 99, 99],
        [18, 21, 26, 66, 99, 99, 99, 99],
        [24, 26, 56, 99, 99, 99, 99, 99],
        [47, 66, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        ])

QN= np.ones((8,8))

##########################################
### Data Set #############################
##########################################

ImgDir = r'.' # Address of folder with files (do nor delete `r``)

Images=[
    {"Filename":"img1.jpg", "ROIs":[
        [1900, 2400, 128, 128],
        [1800, 2200, 128, 128],
        [1400, 1300, 128, 128],
    ]},
    {"Filename":"img2.jpg", "ROIs":[
        [1700, 1700, 128, 128],
        [1900, 1300, 128, 128],
        [2100, 1800, 128, 128],
    ]},
    {"Filename":"img3.jpg", "ROIs":[
        [1700, 1600, 128, 128],
        [1500, 2200, 128, 128],
        [1900, 1700, 128, 128],
    ]},
    {"Filename": "img4.jpg", "ROIs": [
        [1200, 1700, 128, 128],
        [1500, 2200, 128, 128],
        [900, 2600, 128, 128],
    ]},
]


##########################################
### Functions to  ########################
##########################################

# chose JPEG container definition
class JPEG_class:
    Y = np.array([])
    Cb = np.array([])
    Cr = np.array([])
    ChromaRatio = "4:4:4"
    QY = np.ones((8, 8))
    QC = np.ones((8, 8))
    shape = (0, 0, 3)

def CompressJPEG(RGB,Ratio="4:4:4",QY=QN,QC=QN):
    JPEG=JPEG_class()
    JPEG.shape=RGB.shape
    JPEG.ChromaRatio=Ratio
    JPEG.QY=QY
    JPEG.QC=QC
    YCrCb = cv2.cvtColor(RGB, cv2.COLOR_RGB2YCrCb).astype(int)
    Y, Cr, Cb = cv2.split(YCrCb)
    JPEG.Y=Y
    if Ratio == "4:2:2":
        JPEG.Cb = Cb[:, ::2]
        JPEG.Cr = Cr[:, ::2]
    elif Ratio == "4:2:0":
        pass
    else:
        JPEG.Cb = Cb
        JPEG.Cr = Cr
    before_Y = JPEG.Y.size
    before_Cr = JPEG.Cr.size
    before_Cb = JPEG.Cb.size

    JPEG.Y = CompressLayer(JPEG.Y, JPEG.QY)
    JPEG.Cr = CompressLayer(JPEG.Cr, JPEG.QC)
    JPEG.Cb = CompressLayer(JPEG.Cb, JPEG.QC)
    JPEG.Y=byteRun_encode(JPEG.Y)
    JPEG.Cr=byteRun_encode(JPEG.Cr)
    JPEG.Cb=byteRun_encode(JPEG.Cb)
    print("Przed kompresją:")
    print("Y size:", before_Y)
    print("Cr size:", before_Cr)
    print("Cb size:", before_Cb)
    print("Po kompresji:")
    print("Ratio:", JPEG.ChromaRatio)
    if np.array_equal(JPEG.QY, QN) and np.array_equal(JPEG.QC, QN):
        print("Kwantyzacja: False")
    else:
        print("Kwantyzacja: True")
    print("Y size:", len(JPEG.Y))
    print("Cr size:", len(JPEG.Cr))
    print("Cb size:", len(JPEG.Cb))
    print("----------------")
    return JPEG


def DecompressJPEG(JPEG):
    Y = byteRun_decode(JPEG.Y)
    Cr = byteRun_decode(JPEG.Cr)
    Cb = byteRun_decode(JPEG.Cb)
    Y = DecompressLayer(Y, JPEG.QY)
    Cr = DecompressLayer(Cr, JPEG.QC)
    Cb = DecompressLayer(Cb, JPEG.QC)
    if JPEG.ChromaRatio == "4:2:2":
        Cb = np.repeat(Cb, 2, axis=1)
        Cr = np.repeat(Cr, 2, axis=1)
    elif JPEG.ChromaRatio == "4:2:0":
        pass



    YCrCb = np.dstack([Y, Cr, Cb])
    YCrCb = np.clip(YCrCb, 0, 255).astype(np.uint8)

    RGB = cv2.cvtColor(YCrCb, cv2.COLOR_YCrCb2RGB)
    return RGB

def byteRun_encode(data):
    data=data.astype(int)
    x = np.array([len(data.shape)])
    x = np.concatenate([x, data.shape])
    flat=data.flatten()
    n=len(flat)

    buffer=np.zeros((n*2))
    out_idx=0
    i=0
    while i<n:
        start_i = i
        if i<n-1 and flat[i]==flat[i+1]:
            count=1
            for j in range(i+1,n):
                if(flat[j] == flat[i]):
                    count+=1

                else:
                    break
            full_count=count
            while count > 128:
                buffer[out_idx] = -127
                buffer[out_idx + 1] = flat[i]
                out_idx += 2
                count -= 128

            buffer[out_idx]=1-count
            buffer[out_idx+1]=flat[i]
            out_idx+=2
            i+=full_count
        else:
            values=[flat[i]]
            for j in range(i+1,n):
                if j<n-1 and flat[j]==flat[j+1]:
                    break

                values.append(flat[j])
            count=len(values)
            idx=0
            while count > 128:
                buffer[out_idx] = 127
                out_idx += 1
                for k in range(128):
                    buffer[out_idx] = values[idx]
                    out_idx += 1
                    idx += 1
                count -= 128
            buffer[out_idx] = count-1
            out_idx+=1
            for k in range(count):
                buffer[out_idx] = values[idx]
                out_idx += 1
                idx += 1
            i+=len(values)

    compressed_data = buffer[:out_idx]
    result= np.concatenate([x,compressed_data])
    return result.astype(int)

def byteRun_decode(data):
    OGshape = data[1:int(data[0] + 1)]
    encoded = data[int(data[0]) + 1:]
    decoded=[]
    i=0
    while i<len(encoded):
        count=encoded[i]
        if count<0:

            value=encoded[i+1]
            decoded.extend([value] * (1-count))
            i+=2
        else:
            decoded.extend(encoded[i+1:i+1+count+1])
            i += 1+ count+1
    decoded = np.array(decoded)
    return decoded.reshape(OGshape).astype(int)

def zigzag(A):
    template= np.array([
            [0,  1,  5,  6,  14, 15, 27, 28],
            [2,  4,  7,  13, 16, 26, 29, 42],
            [3,  8,  12, 17, 25, 30, 41, 43],
            [9,  11, 18, 24, 31, 40, 44, 53],
            [10, 19, 23, 32, 39, 45, 52, 54],
            [20, 22, 33, 38, 46, 51, 55, 60],
            [21, 34, 37, 47, 50, 56, 59, 61],
            [35, 36, 48, 49, 57, 58, 62, 63],
            ])
    if len(A.shape)==1:
        B=np.zeros((8,8))
        for r in range(0,8):
            for c in range(0,8):
                B[r,c]=A[template[r,c]]
    else:
        B=np.zeros((64,))
        for r in range(0,8):
            for c in range(0,8):
                B[template[r,c]]=A[r,c]
    return B

def dct2(a):
    return scipy.fftpack.dct( scipy.fftpack.dct( a.astype(float), axis=0, norm='ortho' ), axis=1, norm='ortho' )

def idct2(a):
    return scipy.fftpack.idct( scipy.fftpack.idct( a.astype(float), axis=0 , norm='ortho'), axis=1 , norm='ortho')

def CompressBlock(block,Q):
    d=dct2(block)
    qd=np.round(d/Q).astype(int)
    z=zigzag(qd)
    return z

def DecompressBlock(vector, Q):
    block=zigzag(vector)
    block=block*Q
    block=idct2(block)
    return block

def CompressLayer(L,Q):
    S=np.array([])
    for w in range(0,L.shape[0],8):
        for k in range(0,L.shape[1],8):
            block=L[w:(w+8),k:(k+8)]
            S=np.append(S, CompressBlock(block,Q))
    return S
def DecompressLayer(S,Q):
    if S.shape[0] == 128*128:
        L = np.zeros((128, 128))
    else:
        L = np.zeros((128, 64))
    for idx,i in enumerate(range(0,S.shape[0],64)):
        vector=S[i:(i+64)]
        m=L.shape[1]/8
        k=int((idx%m)*8)
        w=int((idx//m)*8)
        L[w:(w+8),k:(k+8)]=DecompressBlock(vector,Q)
    return L
# rest of the functions here 

##########################################
### Main Program  ########################
##########################################

def plot_comparisone(counter, OG,Decomp,figsize=(5,8) ):
    fig, axs = plt.subplots(4, 2 ,num=counter, sharex=True, sharey=True,figsize=figsize )
    # obraz oryginalny 
    axs[0,0].imshow(OG) #RGB 
    PRZED_YCrCb=cv2.cvtColor(OG,cv2.COLOR_RGB2YCrCb)
    axs[1,0].imshow(PRZED_YCrCb[:,:,0],cmap='gray') 
    axs[2,0].imshow(PRZED_YCrCb[:,:,1],cmap='gray')
    axs[3,0].imshow(PRZED_YCrCb[:,:,2],cmap='gray')
    
    axs[0,0].set_title("Oryginał")
    axs[1,0].set_title("Y")
    axs[2,0].set_title("Cr")
    axs[3,0].set_title("Cb")

    # obraz po dekompresji
    axs[0,1].imshow(Decomp) #RGB 
    PO_YCrCb=cv2.cvtColor(Decomp,cv2.COLOR_RGB2YCrCb)
    axs[1,1].imshow(PO_YCrCb[:,:,0],cmap='gray')
    axs[2,1].imshow(PO_YCrCb[:,:,1],cmap='gray')
    axs[3,1].imshow(PO_YCrCb[:,:,2],cmap='gray')
    
    axs[0,1].set_title("Po kompresji")
    axs[1,1].set_title("Y")
    axs[2,1].set_title("Cr")
    axs[3,1].set_title("Cb")
    
    for ax in axs.flatten():
        ax.set_axis_off() 
    
    return fig

if Test:
    img=plt.imread(os.path.join(ImgDir,Images[0]["Filename"]))
    ROI=Images[0]["ROIs"][0]
    fragment= img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2]]
    Counter=1
    for Chroma in Chroma_options:
        for Quant in Quant_options:
            if Quant:
                tJPEG=CompressJPEG(fragment,Ratio=Chroma,QY=QY,QC=QC)
            else:
                tJPEG=CompressJPEG(fragment,Ratio=Chroma,QY=QN,QC=QN)
                
            New_Fragment=DecompressJPEG(tJPEG)
                
            f = plot_comparisone(Counter,fragment,New_Fragment)
            f.suptitle("Plik: {} Chroma: {} Kwantyzacja: {}".format(Images[0]["Filename"],Chroma,Quant))
            Counter+=1
    plt.show()
    
else:
    # generate raport
    document = Document()
    document.add_heading('Report',0) # tworzenie nagłówków druga wartość to poziom nagłówka 
    document.add_paragraph("Autor: Artur Mizula")
    document.add_section()
    document.add_heading("Fragmenty wygenerowane na podstawie działania funkcji",1)
    Counter=1
    for file_dict in Images:
        filename=file_dict['Filename']
        img=plt.imread(os.path.join(ImgDir,filename))
        for ROI in file_dict['ROIs']:
            fragment= img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2]]
            
            for Chroma in Chroma_options:
                for Quant in Quant_options:
                    if Quant:
                        tJPEG=CompressJPEG(fragment,Ratio=Chroma,QY=QY,QC=QC)
                    else:
                        tJPEG=CompressJPEG(fragment,Ratio=Chroma,QY=QN,QC=QN)
                        
                    New_Fragment=DecompressJPEG(tJPEG)
                        
                    f = plot_comparisone(Counter,fragment,New_Fragment)
                    f.suptitle("Plik: {} Chroma: {} Kwantyzacja: {}".format(filename,Chroma,Quant))
                    memfile = BytesIO() 
                    f.savefig(memfile)
                    document.add_picture(memfile, width=Inches(6)) # set document size
                    memfile.close()
                    f.clf()
                    plt.close(f)
    document.add_section()
    document.add_heading("Podsumowanie i wnioski",1)
    document.add_paragraph("W ćwiczeniu wykonano uproszczoną kompresję JPEG dla wybranych fragmentów obrazów o rozmiarze 128x128 pikseli.Użyto kompresji ByteRun."
    "Na podstawie uzyskanych wyników można zauważyć, że największy wpływ na zmniejszenie rozmiaru danych miała kwantyzacja. Użycie kwantyzacji powodował że rozmiar kanału Y zmniejszał się zwykle około 75%. Kanały Cr i Cb po kwantyzacji były kompresowane jeszcze skuteczniej zmniejszały się nawet o ponad 90%."
    "Wariant 4:2:2 dodatkowo zmniejszał ilość danych w kanałach Cr i Cb. Rozmiary kanałów były mniejsze niż przy wariancie 4:4:4. Różnica jakości obrazu była stosunkowo niewielka."
    "Bez kwantyzacji kompresja była znacznie mniej skuteczna. W większości przypadkach rozmiar danych po ByteRun dla kanału Y był zbliżony do rozmiaru początkowego lub nawet większy.Dla kanałów Cb i Cr zmiejsał się około 32%."
    "Najlepszy kompromis między jakością obrazu a stopniem kompresji uzyskano dla wariantu 4:2:2 z włączoną kwantyzacją. Obraz po dekompresji zachowywał najważniejsze cechy oryginału, a jednocześnie rozmiar danych był wyraźnie mniejszy.")
    document.save(OutputRaportFile) 