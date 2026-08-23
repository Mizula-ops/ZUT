import cv2
import numpy as np
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO
from skimage.metrics import structural_similarity as ssim
IMAGE_WIDTH = 640
IMAGE_HEIGHT = 480

CANVAS_WIDTH = 640
CANVAS_HEIGHT = 480
results = []
generated_images = []
resized_images = []

LAYERS_COUNT =3
sizes = [
    (320, 240),
    (640, 480),
    (800, 600),
    (1280, 960),
    (1920, 1440)
]
def draw_line(buffer, x1, y1, x2, y2, color):
    x1=int(round(x1))
    y1=int(round(y1))
    x2=int(round(x2))
    y2=int(round(y2))
    dx = abs(x2-x1)
    dy = abs(y2-y1)
    if x1<x2:
        sx=1
    else:
        sx=-1
    if y1<y2:
        sy=1
    else:
        sy=-1
    err = dx-dy
    height, width = buffer.shape[:2]
    while True:
        if 0 <= x1 < width and 0 <= y1 < height:
            buffer[y1,x1]=color
        if x1==x2 and y1==y2:
            break
        e2=2*err
        if e2> -dy:
            err -=dy
            x1+=sx
        if e2< dx:
            err +=dx
            y1 +=sy
def draw_circle(buffer,cx, cy, radius, color):
    cx=int(round(cx))
    cy=int(round(cy))
    radius=int(round(radius))
    x=0
    y=radius
    p=5/4-radius
    h,w = buffer.shape[:2]

    while x <= y:
        points = [
            (cx + x, cy + y),
            (cx + x, cy - y),
            (cx - x, cy + y),
            (cx - x, cy - y),
            (cx + y, cy + x),
            (cx + y, cy - x),
            (cx - y, cy + x),
            (cx - y, cy - x)
        ]

        for px, py in points:
            if 0 <= px < w and 0 <= py < h:
                buffer[py, px] = color
        if p < 0:
            p = p + 2*x + 1
            x = x + 1
        else:
            p = p + 2*x + 1 - 2*y
            x = x + 1
            y = y - 1
def draw_polygon(buffer, points, color):
    for i in range(len(points)):
        x1, y1 = points[i]
        x2, y2 = points[(i + 1) % len(points)]
        draw_line(buffer, x1, y1, x2, y2, color)


# Draw triangle from 3 points
def draw_triangle(buffer, p1, p2, p3, color):
    draw_line(buffer, p1[0], p1[1], p2[0], p2[1], color)
    draw_line(buffer, p2[0], p2[1], p3[0], p3[1], color)
    draw_line(buffer, p3[0], p3[1], p1[0], p1[1], color)


# Draw rectangle from two corner points (p1, p2)
def draw_rectangle(buffer, p1, p2, color):
    x1, y1 = p1
    x2, y2 = p2

    draw_line(buffer, x1, y1, x2, y1, color)
    draw_line(buffer, x2, y1, x2, y2, color)
    draw_line(buffer, x2, y2, x1, y2, color)
    draw_line(buffer, x1, y2, x1, y1, color)
def fill_from_edges(buffer, color):
    h,w = buffer.shape[:2]
    for y in range(h):
        x=0
        while x<w:
            while x<w and buffer[y,x,0]<0:
                x+=1
            x_start=x
            while x < w and buffer[y, x, 0] >= 0:
                x += 1
            while x < w and buffer[y, x, 0] < 0:
                x += 1
            x_end=x
            if x_start < w and x_end < w:

                for px in range(x_start, x_end + 1):
                    buffer[y, px] = color
            x += 1
def blend(old_color, new_color, alpha):
    alpha = alpha / 255

    return (alpha * np.array(new_color) +(1 - alpha) * np.array(old_color)).astype(np.uint8)

def draw(scene):
    width = scene["Canvas"]["width"]
    height = scene["Canvas"]["height"]
    bg = scene["Canvas"]["background_color"]

    width_scale = width / CANVAS_WIDTH
    height_scale = height / CANVAS_HEIGHT

    def scale_radius(radius):
        return radius * min(width_scale, height_scale)

    image = np.zeros((height, width, 3), dtype=np.uint8)
    image[:] = bg

    shapes = sorted(scene["Shapes"], key=lambda s: s.get("Z_layer", 0))

    for shape in shapes:
        buffer = np.zeros_like(image, dtype=int) - 99
        color = shape["color"]
        alpha = shape.get("alpha", 255)

        if shape["type"] == "line":
            x1, y1 = PixelCoord(*shape["p1"])
            x2, y2 = PixelCoord(*shape["p2"])
            draw_line(buffer, x1, y1, x2, y2, color)

        elif shape["type"] == "circle":
            cx, cy = PixelCoord(*shape["center"])
            radius = scale_radius(shape["radius"])

            draw_circle(buffer, cx, cy, radius, color)
            if shape.get("fill", False):
                fill_from_edges(buffer, color)
            draw_circle(buffer, cx, cy, radius, color)

        elif shape["type"] == "triangle":
            p1 = PixelCoord(*shape["points"][0])
            p2 = PixelCoord(*shape["points"][1])
            p3 = PixelCoord(*shape["points"][2])

            draw_triangle(buffer, p1, p2, p3, color)

            if shape.get("fill", False):
                draw_polygon(buffer, [p1, p2, p3], color)
                fill_from_edges(buffer, color)
                draw_triangle(buffer, p1, p2, p3, color)

        elif shape["type"] == "rectangle":
            p1 = PixelCoord(*shape["p1"])
            p2 = PixelCoord(*shape["p2"])

            draw_rectangle(buffer, p1, p2, color)

            if shape.get("fill", False):
                rect_points = [
                    [p1[0], p1[1]],
                    [p2[0], p1[1]],
                    [p2[0], p2[1]],
                    [p1[0], p2[1]]
                ]
                draw_polygon(buffer, rect_points, color)
                fill_from_edges(buffer, color)
                draw_rectangle(buffer, p1, p2, color)

        elif shape["type"] == "polygon":
            scaled_points = [PixelCoord(*point) for point in shape["points"]]

            draw_polygon(buffer, scaled_points, color)
            if shape.get("fill", False):
                fill_from_edges(buffer, color)
                draw_polygon(buffer, scaled_points, color)

        idx = buffer[:, :, 0] >= 0

        if alpha == 255:
            image[idx] = buffer[idx].astype(np.uint8)
        else:
            for y, x in np.argwhere(idx):
                image[y, x] = blend(image[y, x], buffer[y, x], alpha)

    return image

IMAGE = np.zeros((IMAGE_HEIGHT,IMAGE_WIDTH,LAYERS_COUNT),dtype=np.uint8)
scene = {
    "Canvas": {
        "width": 640,
        "height": 480,
        "background_color": [0, 0, 0]
    },
    "Shapes": [
        # Duży okrąg częściowo przesłonięty trójkątem
        {
            "type": "circle",
            "center": [150, 140],
            "radius": 90,
            "color": [0, 120, 255],
            "fill": True,
            "Z_layer": 0
        },
        {
            "type": "polygon",
            "points": [
                [90, 80],
                [250, 150],
                [120, 250]
            ],
            "color": [255, 0, 0],
            "fill": True,
            "Z_layer": 1
        },

        # Prostokąt z dwoma kwadratami w środku w innym kolorze
        {
            "type": "polygon",
            "points": [
                [330, 60],
                [560, 60],
                [560, 190],
                [330, 190]
            ],
            "color": [0, 180, 80],
            "fill": True,
            "Z_layer": 0
        },
        {
            "type": "polygon",
            "points": [
                [365, 95],
                [420, 95],
                [420, 150],
                [365, 150]
            ],
            "color": [233, 225, 255],
            "fill": True,
            "Z_layer": 1
        },
        {
            "type": "polygon",
            "points": [
                [470, 95],
                [525, 95],
                [525, 150],
                [470, 150]
            ],
            "color": [233, 225, 255],
            "fill": True,
            "Z_layer": 1
        },

        # Wielokąt przedstawiający literę L
        {
            "type": "polygon",
            "points": [
                [80, 300],
                [130, 300],
                [130, 390],
                [230, 390],
                [230, 440],
                [80, 440]
            ],
            "color": [180, 0, 255],
            "fill": True,
            "Z_layer": 0
        },

        # Żółty okrąg wysunięty w połowie za brązowego prostokąta
        {
            "type": "circle",
            "center": [455, 340],
            "radius": 70,
            "color": [255, 255, 0],
            "fill": True,
            "Z_layer": 0
        },
        {
            "type": "polygon",
            "points": [
                [390, 300],
                [580, 300],
                [580, 400],
                [390, 400]
            ],
            "color": [120, 70, 20],
            "fill": True,
            "Z_layer": 1
        },

        # Dodatkowe obiekty dla cyfry albumu = 1
        # Kwadrat na samym dole
        {
            "type": "polygon",
            "points": [
                [500, 50],
                [580, 50],
                [580, 130],
                [500, 130]
            ],
            "color": [255, 105, 180],
            "fill": True,
            "Z_layer": 2
        },

        # Czerwony okrąg nad kwadratem
        {
            "type": "circle",
            "center": [540, 90],
            "radius": 45,
            "color": [255, 0, 0],
            "fill": True,
            "Z_layer": 3
        },

        # Biały trójkąt na samej górze
        {
            "type": "polygon",
            "points": [
                [540, 20],
                [600, 120],
                [480, 120]
            ],
            "color": [255, 255, 255],
            "fill": True,
            "Z_layer": 4
        }
    ]
}

results = []
def mse(img1, img2):

    return np.mean((img1.astype(float) - img2.astype(float)) ** 2)
WIDTH_SCALE = 1.0
HEIGHT_SCALE = 1.0
PixelCoord = lambda x, y: (x * WIDTH_SCALE, y * HEIGHT_SCALE)
reference = draw(scene)

for width, height in sizes:
    WIDTH_SCALE = width / CANVAS_WIDTH
    HEIGHT_SCALE = height / CANVAS_HEIGHT

    PixelCoord = lambda x, y: (x * WIDTH_SCALE, y * HEIGHT_SCALE)
    scene["Canvas"]["width"] = width
    scene["Canvas"]["height"] = height
    img = draw(scene)
    img_resized = cv2.resize(img, (640, 480), interpolation=cv2.INTER_LINEAR)
    mse_value = mse(reference, img_resized)
    ssim_value = ssim(
        reference,
        img_resized,
        channel_axis=2

    )
    results.append([width, height, mse_value, ssim_value])
    generated_images.append((width, height, img))
    resized_images.append((width, height, img_resized))
    plt.figure(figsize=(8, 6))
    plt.imshow(img)
    plt.title(f"Obraz wygenerowany: {width}x{height}")
    plt.axis("off")
    plt.show()

    plt.figure(figsize=(8, 6))
    plt.imshow(img_resized)
    plt.title(f"Po przeskalowaniu do 640x480")
    plt.axis("off")
    plt.show()


print("Rozmiar | MSE | SSIM")
for row in results:
    print(row)

fig, axs = plt.subplots(1, len(generated_images), figsize=(20, 5))
for ax, (width, height, img) in zip(axs, generated_images):
    ax.imshow(img)
    ax.set_title(f"{width}x{height}")
    ax.axis("off")
plt.tight_layout()

buf_generated = BytesIO()
plt.savefig(buf_generated, format="png", dpi=150)
buf_generated.seek(0)
plt.show()
plt.close(fig)


fig, axs = plt.subplots(1, len(resized_images), figsize=(20, 5))
for ax, (width, height, img_resized) in zip(axs, resized_images):
    ax.imshow(img_resized)
    ax.set_title(f"{width}x{height} -> 640x480")
    ax.axis("off")
plt.tight_layout()

buf_resized = BytesIO()
plt.savefig(buf_resized, format="png", dpi=150)
buf_resized.seek(0)
plt.show()
plt.close(fig)


document = Document()
document.add_heading("Sprawozdanie - Rasteryzacja grafiki wektorowej", 0)

document.add_heading("Wygenerowane obrazy", 1)
document.add_picture(buf_generated, width=Inches(6.5))

document.add_heading("Obrazy przeskalowane do 640x480", 1)
document.add_picture(buf_resized, width=Inches(6.5))

document.add_heading("Wyniki porównania", 1)
table = document.add_table(rows=1, cols=4)
table.style = "Table Grid"

headers = table.rows[0].cells
headers[0].text = "Rozdzielczość"
headers[1].text = "Szerokość"
headers[2].text = "MSE"
headers[3].text = "SSIM"

for width, height, mse_value, ssim_value in results:
    row = table.add_row().cells
    row[0].text = f"{width}x{height}"
    row[1].text = str(width)
    row[2].text = f"{mse_value:.4f}"
    row[3].text = f"{ssim_value:.6f}"

document.add_heading("Wnioski", 1)
document.add_paragraph(
    "Dla rozdzielczości 640×480 porównano dwa obrazy wygenerowane w tej samej rozdzielczości "
    "w celu sprawdzenia poprawności działania algorytmu. Uzyskano wartości MSE równe 0 oraz SSIM"
    " równe 1, co oznacza, że obrazy są identyczne. Dla pozostałych rozdzielczości różnice wynikają z "
    "procesu skalowania obrazu do rozmiaru referencyjnego 640×480. Wartości SSIM pozostają wysokie, co oznacza, "
    "że geometria sceny została poprawnie przeskalowana i zachowano wysokie podobieństwo wizualne obrazów."
    "Przy najniższej rozdzielczości 320×240 można zauważyć większą pikselizację obrazu, "
    "szczególnie widoczną na krawędziach okręgów, które tracą swoją gładkość i stają się bardziej schodkowane. "
    "Mój indeks albumu to 55831 czyli moim dodatkowym zadaniem było narysowanie okręgu, kwadratu i trójkąta który jest widoczny w prawym górnym rogu."

)

document.save("report.docx")
print("Zapisano report.docx")
