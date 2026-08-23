import numpy as np
import matplotlib.pyplot as plt
import scipy
import sounddevice as sd
import soundfile as sf

#ZAD1
data, fs =sf.read('SOUND_INTRO/sound1.wav', dtype='float32')

#sd.play(data,fs)
# status = sd.wait()
lewy=data[:,0]
prawy=data[:,1]
mono=data.mean(axis=1)

sf.write('soundL.wav', lewy, fs)
sf.write('soundR.wav', prawy, fs)
sf.write('sound_mix.wav', mono, fs)
plt.subplot(3,1,1)
plt.plot(lewy)
plt.subplot(3,1,2)
plt.plot(prawy)
plt.subplot(3,1,3)
plt.plot(mono)
plt.show()

#WIDMO
data, fs =sf.read('SIN/sin_440Hz.wav', dtype='int32')
#ZAd2
def plotAudio(signal,Fs,TimeMargin=[0,0.02]):
    fsize = 2 ** 8
    yf=scipy.fftpack.fft(signal,fsize)
    signalDB=20*np.log10(np.abs(yf[:fsize//2]))
    fsDb=np.arange(0,Fs/2,Fs/fsize)


    plt.subplot(2,1,1)
    plt.plot(np.arange(0, signal.shape[0]) / Fs, signal)
    plt.xlabel('Seconds')
    plt.ylabel('Amplitude')
    plt.xlim(TimeMargin[0], TimeMargin[1])

    plt.subplot(2,1,2)
    plt.plot(fsDb,signalDB)
    plt.xlabel('Frequency (Hz)')
    plt.ylabel('Amplitude [dB]')
    plt.show()

plotAudio(data,fs)

#ZAD3
def plotAudioModify(signal,Fs,axs,fsize,TimeMargin=[0,0.02]):
    yf=scipy.fftpack.fft(signal,fsize)
    signalDB=20*np.log10(np.abs(yf[:fsize//2])+1e-12)
    fsDb=np.arange(0,Fs/2,Fs/fsize)


    axs[0].plot(np.arange(0, signal.shape[0]) / Fs, signal)
    axs[0].set_xlabel('Seconds')
    axs[0].set_ylabel('Amplitude')
    axs[0].set_xlim(TimeMargin[0], TimeMargin[1])

    axs[1].plot(fsDb,signalDB)
    axs[1].set_xlabel('Frequency (Hz)')
    axs[1].set_ylabel('Amplitude [dB]')

    maxIdx=np.argmax(signalDB)
    maxSignalDb=signalDB[maxIdx]
    maxFsDb=fsDb[maxIdx]
    return maxSignalDb, maxFsDb


from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO

document = Document()
document.add_heading('Zadanie 3', 0)  # tworzenie nagłówków druga wartość to poziom nagłówka

files = ['SIN/sin_60Hz.wav', 'SIN/sin_440Hz.wav', 'SIN/sin_8000Hz.wav']
fsize = [2**8, 2**12,2**16]
for file in files:
    document.add_heading('Plik - {}'.format(file), 2)
    for i, size in enumerate(fsize):
        document.add_heading('Fsize {}'.format(fsize), 3)  # nagłówek sekcji, mozę być poziom wyżej
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))  # tworzenie plota

        data, fs = sf.read(file, dtype='int32')
        maxdB,maxHz = plotAudioModify(data,fs,axs,size)
        fig.suptitle('Fsize {}'.format(size))  # Tytuł wykresu
        fig.tight_layout(pad=1.5)  # poprawa zytelności
        memfile = BytesIO()  # tworzenie bufora
        fig.savefig(memfile)  # z zapis do bufora

        document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

        memfile.close()
        ############################################################
        # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
        document.add_paragraph(f'Max amplitude:{maxdB:.2f} dB, Max frequency: {maxHz:.2f}Hz')
        ############################################################

document.save('report.docx')  # zapis do pliku