import numpy as np
import matplotlib.pyplot as plt
import cv2
#CZESC 1
def imgT0Float(img):
    if (np.issubdtype(img.dtype, np.floating)):
        return img
    if (np.issubdtype(img.dtype, np.integer)):
        return img/255
    return img
def imgToInt8(img):
    if (np.issubdtype(img.dtype, np.integer)):
        return img
    if (np.issubdtype(img.dtype, np.floating)):
        imgInt=img*255
        return imgInt.astype('uint8')
    return img


nazwy=["A1.png","A2.jpg","A3.png","A4.jpg","B01.png","B02.jpg"]
img=[]
for i in nazwy:
    img.append(plt.imread("IMG_INTRO/"+i))
for i in range(len(img)):
    print("Zdjecie:", nazwy[i])
    print(img[i].dtype)
    print(img[i].shape)
    print(np.min(img[i]),np.max(img[i]))
for i in img:
    imgToInt8(i)
#CZESC 2
plt.imshow(img[2])
plt.show()
R=img[2][:,:,0]
plt.imshow(R,cmap=plt.cm.gray,vmin=np.min(R),vmax=np.max(R))
plt.show()
#ZAD2
imgCv2 =cv2.imread("IMG_INTRO/A3.png")
img_RGB = cv2.cvtColor(imgCv2, cv2.COLOR_BGR2RGB)
Y2 = img_RGB[:, :, 0] *0.299 + 0.587* img_RGB[:, :, 1]  + 0.114* img_RGB[:, :, 2]
Y2=imgToInt8(Y2)
cv2.imshow('Grayscale',Y2)
cv2.waitKey(0)
cv2.destroyAllWindows()

#Zadanie Obraz2
def colorf(img):
    tabB1 = []
    for i in range(9):
        tabB1.append(img.copy())
    tabB1[1] = tabB1[1][:, :, 0] * 0.299 + 0.587 * tabB1[1][:, :, 1] + 0.114 * tabB1[1][:, :, 2]
    tabB1[2] = tabB1[2][:, :, 0] * 0.2126 + 0.7152 * tabB1[2][:, :, 1] + 0.0722 * tabB1[2][:, :, 2]
    tabB1[3] = tabB1[3][:, :, 0]
    tabB1[4] = tabB1[4][:, :, 1]
    tabB1[5] = tabB1[5][:, :, 2]
    tabB1[6][:, :, 1] = 0
    tabB1[6][:, :, 2] = 0
    tabB1[7][:, :, 0] = 0
    tabB1[7][:, :, 2] = 0
    tabB1[8][:, :, 0] = 0
    tabB1[8][:, :, 1] = 0
    return tabB1
tabB1=colorf(img[4])


for i in range(9):
    plt.subplot(3,3,i+1)
    if(i<6):
        plt.imshow(tabB1[i],cmap=plt.cm.gray)
    else:
        plt.imshow(tabB1[i])
plt.show()
#Zadanie Obraz 3
import pandas as pd
import cv2
import numpy as np
import matplotlib.pyplot as plt

Fragments=[]
start_w=0
start_k=0
rozmiar=200
for i in range(3):
    for j in range(3):
        w1=start_w+i*rozmiar
        k1=start_k+j*rozmiar
        w2=w1+rozmiar
        k2=k1+rozmiar
        Fragments.append([w1,k1,w2,k2])
df = pd.DataFrame()

df = pd.DataFrame(data={'Filename':['IMG_INTRO/B01.png'],'Grayscale':[False],
                        'Fragments':[Fragments]
                        })

print(df)

from docx import Document
from docx.shared import Inches
from io import BytesIO

document = Document()
document.add_heading('Zadanie Obraz 3',0)

for index, row in df.iterrows():
    img = plt.imread(row['Filename'])
    if row['Grayscale']:
        pass
    else:
        pass
        # Obraz kolowowy
    if row['Fragments'] is not None:
        # mamy nie pustą listę fragmentów
        for f in row['Fragments']:
            fragment = img[f[0]:f[2],f[1]:f[3]].copy()
            tabB1=colorf(fragment)
            document.add_heading(f'Fragment {f[0]}:{f[2]} - {f[1]}:{f[3]}')
            for i in range(9):
                plt.subplot(3, 3, i + 1)
                if (i < 6):
                    plt.imshow(tabB1[i], cmap=plt.cm.gray)
                else:
                    plt.imshow(tabB1[i])

            memfile = BytesIO()
            plt.savefig(memfile, format='png')
            memfile.seek(0)

            document.add_picture(memfile, width=Inches(6))

            memfile.close()

document.save('report.docx')
