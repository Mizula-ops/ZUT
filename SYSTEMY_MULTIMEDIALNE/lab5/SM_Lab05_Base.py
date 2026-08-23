import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import cv2
import os
from scipy.interpolate import interp1d
import scipy
from docx import Document
from docx.shared import Inches
from io import BytesIO
import soundfile as sf

##########################################
### Settings #############################
##########################################

Test=False
Kwant_Test=False # test Kwant function

Interpolation_kind=["cubic","linear"] # what flag for switching interpolation

PlotSettings={
    "Bits":[4,8,16,24],
    "Decimation":[2,4,6,10,24],
    "InterpolationFrequency":[2000,4000, 8000,11999, 16000, 16953, 24000, 41000]
}
ListeningSettings={
    "Bits":[4,8],
    "Decimation":[4,6,10,24],
    "InterpolationFrequency":[4000, 8000,11999, 16000, 16953]
}

OutputRaportFile = "report.docx"
OutputFolder="NewSIN" # place for all your new audio files will be

##########################################
### Data Set #############################
##########################################

AudioDir = r'.' # Address of folder with files (do nor delete `r``)


SinFiles=[
    {"File":"SIN/sin_60Hz.wav","TimeMargin":[0,0.02]},
    {"File":"SIN/sin_440Hz.wav","TimeMargin":[0,0.005]},
    {"File":"SIN/sin_8000Hz.wav","TimeMargin":[0,0.000250 ]},
    {"File":"SIN/sin_combined.wav","TimeMargin":[0,0.0002]},
    ] # list of dicts with file names with Sinus singals and fragment that will be displayed
SingFiles=["SING/sing_high1.wav","SING/sing_low1.wav","SING/sing_medium1.wav"] # list of file names of with Singing Voice

##########################################
### Functions to  ########################
##########################################



def plotAudio(Signal,Fs,axs,fsize,TimeMargin=[0,0.02]):
    yf = scipy.fftpack.fft(Signal, fsize)
    signalDB = 20 * np.log10(np.abs(yf[:fsize // 2])+1e-10)
    fsDb = np.arange(0, Fs / 2, Fs / fsize)


    axs[0].plot(np.arange(0, Signal.shape[0]) / Fs, Signal)
    axs[0].set_xlabel('Seconds')
    axs[0].set_ylabel('Amplitude')
    axs[0].set_xlim(TimeMargin[0], TimeMargin[1])

    axs[1].plot(fsDb, signalDB)
    axs[1].set_xlabel('Frequency (Hz)')
    axs[1].set_ylabel('Amplitude [dB]')


def Kwant(data,bit):
    d=(2**bit)-1
    if np.issubdtype(data.dtype,np.floating):
        m=-1
        n=1
    else:
        m = np.iinfo(data.dtype).min
        n = np.iinfo(data.dtype).max
    # Przygotowanie sygnalu
    DataF = data.astype(float)
    DataF = (DataF - m)/(n-m)
    # Kwantyzacja
    DataF = (np.round(DataF * d)) / d
    # Powrot do orginalnej przestrzeni
    DataF=(DataF*(n-m))+m
    return DataF.astype(data.dtype)

def decimation(Signal,Fs,step):
    NewSignal=Signal[::step].copy()
    NewFs=Fs//step
    return NewSignal,NewFs

def interpolation(Signal,Fs,NewFs,kind):
    N=len(Signal)
    N1=int(N*NewFs/Fs)
    x=np.linspace(0,N/Fs,N)
    x1 = np.linspace(0,N/Fs,N1)
    y=0
    if kind == "linear":
        metode_lin = interp1d(x, Signal)
        y = metode_lin(x1).astype(Signal.dtype)
    else:
        metode_nonlin=interp1d(x,Signal,kind=kind)
        y=metode_nonlin(x1).astype(Signal.dtype)
    return y


##########################################
### Main Program  ########################
##########################################

if Test:
    counter=1
    if Kwant_Test:
        T_X=[
            np.round(np.linspace(0,255,255,dtype=np.uint8)),
            np.round(np.linspace(np.iinfo(np.int32).min,np.iinfo(np.int32).max,1000,dtype=np.int32)),
            np.linspace(-1,1,10000),
        ]
        Bits=[1,2,4]
        for X in T_X:
            for bit in Bits:
                kwanted=Kwant(X,bit)
                print(f"Bits {bit} == {2**bit} values, unique values {np.unique(kwanted).size}. Dtype before {X.dtype} and after {kwanted.dtype}")
                plt.figure(counter)
                plt.plot(X,kwanted)
                plt.title(f"{bit} bit")
                counter+=1
                
    else:
        file=SinFiles[0]
        Signal, Fs = sf.read(os.path.join(AudioDir,file["File"]), dtype='float32') 
        # test decimation
        dec_Signal,dec_Fs=decimation(Signal,Fs,10)
        f,axs=plt.subplots(2,1,num=counter,figsize=(5,5)) 
        counter+=1
        plotAudio(Signal=dec_Signal,Fs=dec_Fs,axs=axs,fsize=2**12,TimeMargin=file["TimeMargin"])
        f.suptitle(f"{file['File']} Decimation step 10")
        # test interpolation
        for kind in Interpolation_kind:
            Int_Fs=16000
            Int_Signal=interpolation(Signal=Signal,Fs=Fs,NewFs=Int_Fs,kind=kind)
            f,axs=plt.subplots(2,1,num=counter,figsize=(5,5)) 
            counter+=1
            plotAudio(Signal=Int_Signal,Fs=Int_Fs,axs=axs,fsize=2**12,TimeMargin=file["TimeMargin"])
            f.suptitle(f"{file['File']} Interpolation {kind}")
        
        
    plt.show()
    
else:
    # generate raport
    document = Document()
    document.add_heading('Report',0) # tworzenie nagłówków druga wartość to poziom nagłówka 
    document.add_paragraph("Autor: Artur Mizuła")
    document.add_section()
    document.add_heading("Sprawdzanie działania napisanych funkcji na podstawie wykresów",1)
    counter = 1 
    document.add_heading("Testowanie funkcji kwantyzującej",2)
    for file in SinFiles:
        Signal, Fs = sf.read(os.path.join(AudioDir,file["File"]), dtype='float32') 
        for bit in PlotSettings["Bits"]:
            kSignal=Kwant(Signal,bit)
            f,axs=plt.subplots(2,1,num=counter,figsize=(5,5))
            counter += 1
            plotAudio(Signal=kSignal,Fs=Fs,axs=axs,fsize=2**12,TimeMargin=file["TimeMargin"])
            f.suptitle(f"{file['File']} Kwantyzacja {bit}-bitów")
            memfile = BytesIO() 
            f.savefig(memfile)
            document.add_picture(memfile, width=Inches(6)) # set document size
            memfile.close()
            f.clf()
            plt.close(f)
    document.add_heading("Testowanie funkcji decymującej",2)        
    for file in SinFiles:
        Signal, Fs = sf.read(os.path.join(AudioDir,file["File"]), dtype='float32') 
        for step in PlotSettings["Decimation"]:
            dec_Signal,dec_Fs=decimation(Signal,Fs,step)
            f,axs=plt.subplots(2,1,num=counter,figsize=(5,5))
            counter += 1
            plotAudio(Signal=dec_Signal,Fs=dec_Fs,axs=axs,fsize=2**12,TimeMargin=file["TimeMargin"])
            f.suptitle(f"{file['File']} Decimation step {step}")
            memfile = BytesIO() 
            f.savefig(memfile)
            document.add_picture(memfile, width=Inches(6)) # set document size
            memfile.close()
            f.clf()
            plt.close(f)
    document.add_heading("Testowanie funkcji interpolujących",2)        
    for file in SinFiles:
        Signal, Fs = sf.read(os.path.join(AudioDir,file["File"]), dtype='float32') 
        for Int_Fs in PlotSettings["InterpolationFrequency"]:
            for kind in Interpolation_kind:
                Int_Signal=interpolation(Signal=Signal,Fs=Fs,NewFs=Int_Fs,kind=kind)
                f,axs=plt.subplots(2,1,num=counter,figsize=(5,5))
                counter += 1
                handle=plotAudio(Signal=Int_Signal,Fs=Int_Fs,axs=axs,fsize=2**12,TimeMargin=file["TimeMargin"])
                f.suptitle(f"{file['File']} Interpolation {kind} Fs {Int_Fs}")
                memfile = BytesIO() 
                f.savefig(memfile)
                document.add_picture(memfile, width=Inches(6)) # set document size
                memfile.close()
                f.clf()
                plt.close(f)
                document.add_paragraph(f"Tu proszę dostosować do obsługi wyjścia waszej funkcji plot data {handle}")

    document.add_heading("Podsumowanie pierwszej części zadania", 2)

    document.add_paragraph("Kwantyzacja:")
    document.add_paragraph(
        "Im mniejsza liczba bitów, tym mniej poziomów amplitudy przyjmuje sygnał, przez co sygnał 4-bitowy był wyraźnie schodkowy. "
        "Dodatkowo przy małej liczbie bitów w widmie pojawiał się wyraźny szum kwantyzacji. Dla 8 bitów jakość była znacznie lepsza. "
        "Dla 16 i 24 bitów przebieg czasowy był praktycznie nieodróżnialny od oryginalnego, a widmo pozostawało bardzo podobne. "
        "W przypadku sygnałów prostych większa liczba bitów pozwalała na bardzo dobre odwzorowanie sygnału. "
        "Natomiast dla sygnałów złożonych widmo było bardziej zaszumione. "
        "Można zauważyć, że kwantyzacja wpływa głównie na dokładność amplitudy sygnału, a nie na jego częstotliwość. "
        "Zwiększenie liczby bitów zmniejsza poziom szumu kwantyzacji."
    )

    document.add_paragraph("Decymacja:")
    document.add_paragraph(
        "Przy zastosowaniu decymacji z małą wartością kroku przebieg czasowy sygnałów prostych był nadal poprawnie odwzorowany, "
        "a widmo pozostawało czytelne. Wraz ze wzrostem kroku decymacji częstotliwość próbkowania malała, co prowadziło do aliasingu. "
        "Objawiało się to zniekształceniem widma oraz zmianą pozornych częstotliwości sygnału. "
        "Bardzo wyraźny efekt zniekształceń był widoczny dla sygnału sin_8000Hz.wav. "
        "Dla większych kroków sygnał w niektórych przypadkach przyjmował zupełnie inne częstotliwości, a nawet postać sygnału stałego. "
        "W przypadku sygnału złożonego (sin_combined) decymacja nie pozwalała na poprawne odtworzenie ani przebiegu czasowego, ani widma niezależnie od wybranego kroku."
    )

    document.add_paragraph("Interpolacja:")
    document.add_paragraph(
        "Interpolacja działała poprawnie dla sygnałów, gdy częstotliwość próbkowania była odpowiednio duża. "
        "Metoda cubic lepiej odwzorowywała kształt sygnału niż interpolacja liniowa. "
        "Im większa nowa częstotliwość próbkowania (Fs), tym lepiej odwzorowany był przebieg sygnału. "
        "Dla małych wartości Fs interpolacja nie była w stanie poprawnie odtworzyć kształtu sygnału i często prowadziła do uproszczonych przebiegów, "
        "np. odcinków liniowych. "
        "Dla sygnałów złożonych interpolacja nie radziła sobie dobrze przy niskich częstotliwościach próbkowania. "
        "Dopiero dla bardzo dużych wartości Fs przebieg zaczynał w miarę przypominać sygnał oryginalny, jednak nadal był wyraźnie mniej dokładny niż sygnał po kwantyzacji przy większej liczbie bitów."

    )
    document.add_heading("Obserwacje na podstawie odsłuchanych plików", 1)

    document.add_paragraph(
        "W przypadku kwantyzacji dla sygnałów o średnim i wysokim tonie jakość była dobra i porównywalna z oryginałem (dla większej liczby bitów). "
        "Natomiast przy małej liczbie bitów pojawiał się wyraźny szum oraz efekt przesterowania. Kwantyzacja znacznie gorzej radziła sobie "
        "z dźwiękami o niskim tonie – sygnał stawał się zniekształcony, występował szum i był nienaturalnie głośny. "

        "Jeśli chodzi o interpolację oraz decymację, jakość dźwięku oraz występujące defekty były do siebie zbliżone. "
        "Dla sygnałów o średnich i wysokich tonach, przy zbyt małej częstotliwości próbkowania (interpolacja) "
        "lub dużym kroku (decymacja), dźwięk stawał się nienaturalny i przypominał efekt metaliczny (Brzmiało jak UFO). "

        "Natomiast dla dźwięków o niskim tonie interpolacja i decymacja radziły sobie lepiej niż kwantyzacja przy małej liczbie bitów. "
        "Przy nieodpowiednich wartościach parametrów zarówno interpolacja, jak i decymacja powodowały zniekształcenia, "
        "jednak były one zazwyczaj mniejsze niż w przypadku bardzo niskiej rozdzielczości kwantyzacji, gdzie dominował szum."
    )
    for file in SingFiles:
        Signal, Fs = sf.read(os.path.join(AudioDir,file), dtype='float32') 
        sfile=file.split(os.sep)[-1].split('.')
        for bit in ListeningSettings["Bits"]:
            kSignal=Kwant(Signal,bit)
            nfile=f"{sfile[0]}_kwant_{bit}.wav"
            sf.write(os.path.join(OutputFolder,nfile),data=kSignal,samplerate=Fs)
        for step in ListeningSettings["Decimation"]:
            dec_Signal,dec_Fs=decimation(Signal,Fs,step)
            
            nfile=f"{sfile[0]}_dec_{step}.wav"
            sf.write(os.path.join(OutputFolder,nfile),data=dec_Signal,samplerate=dec_Fs)
            
        for Int_Fs in ListeningSettings["InterpolationFrequency"]:
            for kind in Interpolation_kind:
                Int_Signal=interpolation(Signal=Signal,Fs=Fs,NewFs=Int_Fs,kind=kind)
                
                nfile=f"{sfile[0]}_interp_{kind}_{Int_Fs}.wav"
                sf.write(os.path.join(OutputFolder,nfile),data=Int_Signal,samplerate=Int_Fs)
        
        

    document.add_heading("Wnioski",1)

    document.add_paragraph(
        "Na podstawie przeprowadzonych eksperymentów można stwierdzić, że każda z analizowanych operacji wpływa na sygnał w inny sposób i jej skuteczność zależy od rodzaju sygnału oraz dobranych parametrów. "
        "Kwantyzacja przede wszystkim wpływa na dokładność odwzorowania amplitudy sygnału.Przy małej liczbie bitów pojawia się wyraźny szum "
        "kwantyzacji oraz efekt przesterowania. ""Dla sygnałów o niskim tonie zniekształcenia są bardziej zauważalne, "
        "natomiast dla sygnałów o średnim i wysokim tonie przy większej liczbie bitów jakość jest bardzo dobra."
        "Decymacja oraz interpolacja wpływają głównie na odwzorowanie w dziedzinie czasu i częstotliwości. Przy nieodpowiednich parametrach (zbyt duży krok decymacji lub zbyt mała częstotliwość próbkowania przy interpolacji) pojawia się aliasing oraz zniekształcenia sygnału. "
        "Dla sygnałów o niskim tonie metody te radzą sobie lepiej niż kwantyzacja przy niskiej rozdzielczości, natomiast dla tonów średnich i wysokich łatwo pojawiają się nienaturalne efekty dźwiękowe. "
        "W przypadku sygnałów złożonych (np. sin_combined) wszystkie metody mają większe trudności z poprawnym odwzorowaniem sygnału. Szczególnie widoczne są zniekształcenia widma oraz pogorszenie jakości dźwięku. "
        "Najlepsze rezultaty uzyskuje się przy odpowiednim doborze parametrów: dużej liczbie bitów w kwantyzacji."
    )
    document.save(OutputRaportFile)