import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import cv2
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO


##########################################
### Settings #############################
##########################################

Test = False
Scaling_test = False # run only artificial test for scaling methods

ScalesUp = [200,300,500] # list of parameters values
ScalesDown =[5,25,50] # list of parameters values

OutputRaportFile = "raport.docx"

##########################################
### Data Set #############################
##########################################

ImgDir = r'.' # Address of folder with files (do nor delete `r``)


SmallImages=["IMG_SMALL/SMALL_0001.tif", "IMG_SMALL/SMALL_0002.png",
             "IMG_SMALL/SMALL_0003.png", "IMG_SMALL/SMALL_0004.jpg",
             "IMG_SMALL/SMALL_0005.jpg", "IMG_SMALL/SMALL_0006.jpg",
             "IMG_SMALL/SMALL_0007.jpg", "IMG_SMALL/SMALL_0008.jpg",
             "IMG_SMALL/SMALL_0009.jpg", "IMG_SMALL/SMALL_0010.jpg",]

BigImages=[
{
        "Filename": "IMG_BIG/BIG_0001.jpg",
        "ROIs": [
            [1800, 1500, 1200, 1200]
        ]
    },
    {
        "Filename": "IMG_BIG/BIG_0002.jpg",
        "ROIs": [
            [2200, 300, 1200, 1200]
        ]
    },
{
        "Filename": "IMG_BIG/BIG_0003.jpg",
        "ROIs": [
            [450, 2550, 1200, 1200],
        ]
    },
{
        "Filename": "IMG_BIG/BIG_0004.png",
        "ROIs": [
            [600, 200, 700, 700]
        ]
    }
]


##########################################
### Functions to  ########################
##########################################

# Scaling methods

def NearestNeigbourScaling(In_img,scale):
    height=In_img.shape[0]
    width=In_img.shape[1]
    new_height=np.ceil(height*scale/100).astype(int)
    new_width=np.ceil(width*scale/100).astype(int)
    Y=np.linspace(0,height-1,new_height)
    X = np.linspace(0, width - 1, new_width)
    if(len(In_img.shape)<3):
        Out_img=np.zeros((new_height,new_width))
    else:
        Out_img=np.zeros((new_height,new_width,In_img.shape[2]))
    for iy,y in enumerate(Y):
        for ix,x in enumerate(X):
            xx=np.round(x).astype(int)
            yy=np.round(y).astype(int)
            Out_img[iy,ix]=In_img[yy,xx]
    return Out_img.astype(In_img.dtype)

def BilinearScaling(In_img,scale):
    height = In_img.shape[0]
    width = In_img.shape[1]
    new_height = np.ceil(height * scale / 100).astype(int)
    new_width = np.ceil(width * scale / 100).astype(int)
    Y = np.linspace(0, height - 1, new_height)
    X = np.linspace(0, width - 1, new_width)
    if (len(In_img.shape) < 3):
        Out_img = np.zeros((new_height, new_width))
    else:
        Out_img = np.zeros((new_height, new_width, In_img.shape[2]))
    for iy, y in enumerate(Y):
        for ix, x in enumerate(X):
            x1 = int(np.floor(x))
            x2 = int(np.ceil(x))
            y1 = int(np.floor(y))
            y2 = int(np.ceil(y))

            dx = x - x1
            dy = y - y1

            p1 = In_img[y1, x1]
            p2 = In_img[y1, x2]
            p3 = In_img[y2, x1]
            p4 = In_img[y2, x2]

            Out_img[iy, ix] = (p1 *(1 - dx) *(1 - dy) +p2 *dx *(1 - dy) +p3 *(1 - dx) * dy +p4 *dx *dy)
    return Out_img.astype(In_img.dtype)

# Shrinking methods
 
def MeanResizing(In_img,scale):
    height = In_img.shape[0]
    width = In_img.shape[1]
    new_height = np.ceil(height * scale / 100).astype(int)
    new_width = np.ceil(width * scale / 100).astype(int)
    Y = np.linspace(0, height - 1, new_height)
    X = np.linspace(0, width - 1, new_width)
    if (len(In_img.shape) < 3):
        Out_img = np.zeros((new_height, new_width))
    else:
        Out_img = np.zeros((new_height, new_width, In_img.shape[2]))
    for iy_out, yy in enumerate(Y):
        if iy_out > 0:
            y1 = -(yy - Y[iy_out - 1]) / 2
        else:
            y1 = 0
        if iy_out < len(Y) - 1:
            y2 = (Y[iy_out + 1] - yy) / 2 + 1
        else:
            y2 = 0
        iy = np.round(yy + np.arange(y1, y2)).astype(int)
        iy=iy.clip(0,height-1)
        for ix_out, xx in enumerate(X):
            if ix_out > 0:
                x1 = -(xx - X[ix_out - 1]) / 2
            else:
                x1 = 0
            if ix_out < len(X) - 1:
                x2 = (X[ix_out + 1] - xx) / 2 + 1
            else:
                x2 = 0
            ix = np.round(xx + np.arange(x1, x2)).astype(int)
            ix = ix.clip(0, width - 1)
            if len(In_img.shape) < 3:
                fragment = In_img[iy[0]:iy[-1], ix[0]:ix[-1]]
                Out_img[iy_out, ix_out] = np.mean(fragment)
            else:
                fragment = In_img[iy[0]:iy[-1], ix[0]:ix[-1], :]
                Out_img[iy_out, ix_out] = np.mean(fragment, axis=(0, 1))
    return Out_img.astype(In_img.dtype)

def WeightedMeanResizing(In_img,scale):
    height = In_img.shape[0]
    width = In_img.shape[1]
    new_height = np.ceil(height * scale / 100).astype(int)
    new_width = np.ceil(width * scale / 100).astype(int)
    Y = np.linspace(0, height - 1, new_height)
    X = np.linspace(0, width - 1, new_width)
    if (len(In_img.shape) < 3):
        Out_img = np.zeros((new_height, new_width))
    else:
        Out_img = np.zeros((new_height, new_width, In_img.shape[2]))
    for iy_out, yy in enumerate(Y):
        if iy_out > 0:
            y1 = -(yy - Y[iy_out - 1]) / 2
        else:
            y1 = 0
        if iy_out < len(Y) - 1:
            y2 = (Y[iy_out + 1] - yy) / 2 + 1
        else:
            y2 = 0
        iy = np.round(yy + np.arange(y1, y2)).astype(int)
        iy = iy.clip(0, height - 1)
        for ix_out, xx in enumerate(X):
            if ix_out > 0:
                x1 = -(xx - X[ix_out - 1]) / 2
            else:
                x1 = 0
            if ix_out < len(X) - 1:
                x2 = (X[ix_out + 1] - xx) / 2 + 1
            else:
                x2 = 0
            ix = np.round(xx + np.arange(x1, x2)).astype(int)
            ix = ix.clip(0, width - 1)


            if len(In_img.shape) < 3:
                fragment = In_img[iy[0]:iy[-1], ix[0]:ix[-1]]
                h, w = fragment.shape
                sum_val=0
            else:
                fragment = In_img[iy[0]:iy[-1], ix[0]:ix[-1], :]
                h,w = fragment.shape[:2]
                sum_val = np.zeros(fragment.shape[2])

            sum_w = 0
            cx = h // 2
            cy = w // 2
            for a in range(h):
                for b in range(w):
                    dist = abs(a - cx) + abs(b - cy)
                    weight = 1 / (1 + dist)
                    sum_val += fragment[a, b] * weight
                    sum_w += weight
            Out_img[iy_out, ix_out] = sum_val / sum_w

    return Out_img.astype(In_img.dtype)

def MedianResizing(In_img,scale):
    height = In_img.shape[0]
    width = In_img.shape[1]
    new_height = np.ceil(height * scale / 100).astype(int)
    new_width = np.ceil(width * scale / 100).astype(int)
    Y = np.linspace(0, height - 1, new_height)
    X = np.linspace(0, width - 1, new_width)
    if (len(In_img.shape) < 3):
        Out_img = np.zeros((new_height, new_width))
    else:
        Out_img = np.zeros((new_height, new_width, In_img.shape[2]))
    for iy_out, yy in enumerate(Y):
        if iy_out > 0:
            y1 = -(yy - Y[iy_out - 1]) / 2
        else:
            y1 = 0
        if iy_out < len(Y) - 1:
            y2 = (Y[iy_out + 1] - yy) / 2 + 1
        else:
            y2 = 0
        iy = np.round(yy + np.arange(y1, y2)).astype(int)
        iy = iy.clip(0, height - 1)
        for ix_out, xx in enumerate(X):
            if ix_out > 0:
                x1 = -(xx - X[ix_out - 1]) / 2
            else:
                x1 = 0
            if ix_out < len(X) - 1:
                x2 = (X[ix_out + 1] - xx) / 2 + 1
            else:
                x2 = 0
            ix = np.round(xx + np.arange(x1, x2)).astype(int)
            ix = ix.clip(0, width - 1)
            if len(In_img.shape) < 3:
                fragment = In_img[iy[0]:iy[-1], ix[0]:ix[-1]]
                Out_img[iy_out, ix_out] = np.median(fragment)
            else:
                fragment = In_img[iy[0]:iy[-1], ix[0]:ix[-1], :]
                Out_img[iy_out, ix_out] = np.median(fragment, axis=(0, 1))
    ####
    return Out_img.astype(In_img.dtype)

def imgToInt8(img):
    if np.issubdtype(img.dtype, np.floating):
        img = (img * 255).astype(np.uint8)
    return img

def EdgeDetection(img):
    img_uint8 = imgToInt8(img)

    if (len(img.shape) < 3):
        img_cny = img_uint8
        edges = cv2.Canny(img_cny, 100,200)
    else:
        img_uint8 = img_uint8[:, :, :3]
        B = img_uint8[:, :, 0]
        G = img_uint8[:, :, 1]
        R = img_uint8[:, :, 2]
        B_cny = cv2.Canny(B, 50, 200)
        G_cny = cv2.Canny(G, 50, 200)
        R_cny = cv2.Canny(R, 50, 200)
        edges = cv2.merge([B_cny, G_cny, R_cny])

    return edges

##########################################
### Main Program  ########################
##########################################

def plot_resize(img, scale, nnscale, bscale, ed_img, ed_nnscale, ed_bscale, mr_img, wmr_img, mdr_img, oROI, filename,counter,figsize=(5,5)):
    ROI_org = np.array(oROI).astype(int)
    ROI = np.ceil(np.array(oROI) * scale / 100).astype(int)
    f,axs=plt.subplots(4,3,num=counter,figsize=figsize) 
    f.suptitle(f"{filename} ROI: {ROI_org.tolist()}")
    axs[0,0].imshow(img[ROI_org[1]:ROI_org[1]+ROI_org[3],ROI_org[0]:ROI_org[0]+ROI_org[2],:])
    axs[0,0].set_title("Original")
    axs[0,0].set_axis_off()

    axs[0,1].imshow(nnscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[0,1].set_title(f"NN scale {scale}")
    axs[0,1].set_axis_off()

    axs[0,2].imshow(bscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[0,2].set_title(f"Blinear scale {scale}")
    axs[0,2].set_axis_off()

    axs[1,0].imshow(ed_img[ROI_org[1]:ROI_org[1]+ROI_org[3],ROI_org[0]:ROI_org[0]+ROI_org[2],:])
    axs[1,0].set_title("Edges Original")
    axs[1,0].set_axis_off()

    axs[1,1].imshow(ed_nnscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[1,1].set_title("Edges NN")
    axs[1,1].set_axis_off()

    axs[1,2].imshow(ed_bscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[1,2].set_title("Edges Bilinear")
    axs[1,2].set_axis_off()

    axs[2,0].imshow(mr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[2,0].set_title(f"Mean Resizing scale {scale}")
    axs[2,0].set_axis_off()

    axs[2,1].imshow(wmr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[2,1].set_title(f"Weighted Mean Resizing scale {scale}")
    axs[2,1].set_axis_off()

    axs[2,2].imshow(mdr_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[2,2].set_title(f"Median Resizing scale {scale}")
    axs[2,2].set_axis_off()

    axs[3,0].imshow(ed_img[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[3,0].set_title("Edges Mean Resizing")
    axs[3,0].set_axis_off()

    axs[3,1].imshow(ed_nnscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[3,1].set_title("Edges Weighted Mean Resizing")
    axs[3,1].set_axis_off()

    axs[3,2].imshow(ed_bscale[ROI[1]:ROI[1]+ROI[3],ROI[0]:ROI[0]+ROI[2],:])
    axs[3,2].set_title("Edges Median Resizing")
    axs[3,2].set_axis_off()
    return f

def plot_scaling(img, scale, nnscale, bscale, counter, ed_img, ed_nnscale, ed_bscale, file,figsize=(5,5)):
    f,axs=plt.subplots(2,3,num=counter,figsize=figsize) 
    f.suptitle(f"{file}")

    axs[0,0].imshow(img)
    axs[0,0].set_title("Original")
    axs[0,0].set_axis_off()

    axs[0,1].imshow(nnscale)
    axs[0,1].set_title(f"NN scale {scale}")
    axs[0,1].set_axis_off()

    axs[0,2].imshow(bscale)
    axs[0,2].set_title(f"Blinear scale {scale}")
    axs[0,2].set_axis_off()

    axs[1,0].imshow(ed_img)
    axs[1,0].set_title("Edges Original")
    axs[1,0].set_axis_off()

    axs[1,1].imshow(ed_nnscale)
    axs[1,1].set_title("Edges NN")
    axs[1,1].set_axis_off()

    axs[1,2].imshow(ed_bscale)
    axs[1,2].set_title("Edges Bilinear")
    axs[1,2].set_axis_off()
    return f

if Test:
    # test case
    if Scaling_test:
        img= np.zeros((3,3,3),dtype=np.float32)
        img[1,1,:]=1.0
        for scale in ScalesUp:
            f,axs=plt.subplots(1,2)
            nnscale=NearestNeigbourScaling(img,scale)
            axs[0].imshow(nnscale)
            bscale=BilinearScaling(img,scale)
            axs[1].imshow(bscale)
        
    else:
        counter=1
        for scale in ScalesUp:
            img=plt.imread(os.path.join(ImgDir,SmallImages[0]))
            nnscale=NearestNeigbourScaling(img,scale)
            bscale=BilinearScaling(img,scale)
            ed_img=EdgeDetection(img)
            ed_nnscale=EdgeDetection(nnscale)
            ed_bscale=EdgeDetection(bscale)

            f = plot_scaling(img, scale, nnscale, bscale, counter, ed_img, ed_nnscale, ed_bscale, SmallImages[0])

            counter+=1

        for scale in ScalesDown:    
            img=plt.imread(os.path.join(ImgDir,BigImages[0]["Filename"]))

            nnscale=NearestNeigbourScaling(img,scale)
            bscale=BilinearScaling(img,scale)

            ed_img=EdgeDetection(img)
            ed_nnscale=EdgeDetection(nnscale)
            ed_bscale=EdgeDetection(bscale)

            mr_img=MeanResizing(img,scale)
            wmr_img=WeightedMeanResizing(img,scale)
            mdr_img=MedianResizing(img,scale)

            ed_mr_img=EdgeDetection(mr_img)
            ed_wmr_img=EdgeDetection(wmr_img)
            ed_mdr_img=EdgeDetection(mdr_img)
  
            f= plot_resize(img, scale, nnscale, bscale, ed_img, ed_nnscale, ed_bscale, mr_img, wmr_img, mdr_img, BigImages[0]["ROIs"][0], BigImages[0]["Filename"],counter=counter)
            counter+=1
        
    plt.show()
else: 
    # generate raport
    document = Document()
    document.add_heading('Report',0) # tworzenie nagłówków druga wartość to poziom nagłówka 
    document.add_paragraph("Autor: Artur Mizuła ")
    document.add_paragraph("Proszę wstawić mi 2 jeżeli tego nie wyedytuję")
    document.add_section()
    document.add_heading("Test algorytmów powiększania",1)
    counter = 1 
    for file in SmallImages:
        img=plt.imread(os.path.join(ImgDir,file))
        for scale in ScalesUp:
            nnscale=NearestNeigbourScaling(img,scale)
            bscale=BilinearScaling(img,scale)
            ed_img=EdgeDetection(img)
            ed_nnscale=EdgeDetection(nnscale)
            ed_bscale=EdgeDetection(bscale)
            
            f = plot_scaling(img, scale, nnscale, bscale, counter, ed_img, ed_nnscale, ed_bscale, file) # set figszie

            memfile = BytesIO() 
            f.savefig(memfile)
            document.add_picture(memfile, width=Inches(6)) # set document size
            memfile.close()
            f.clf()
    document.add_section()
    document.add_heading("Test algorytmów pomniejszania",1)
    for file_dict in BigImages:
        filename=file_dict['Filename']
        img=plt.imread(os.path.join(ImgDir,filename))
        for scale in ScalesDown:
            nnscale=NearestNeigbourScaling(img,scale)
            bscale=BilinearScaling(img,scale)

            ed_img=EdgeDetection(img)
            ed_nnscale=EdgeDetection(nnscale)
            ed_bscale=EdgeDetection(bscale)

            mr_img=MeanResizing(img,scale)
            wmr_img=WeightedMeanResizing(img,scale)
            mdr_img=MedianResizing(img,scale)

            ed_mr_img=EdgeDetection(mr_img)
            ed_wmr_img=EdgeDetection(wmr_img)
            ed_mdr_img=EdgeDetection(mdr_img)

            for ROI in file_dict['ROIs']:

                f = plot_resize(img, scale, nnscale, bscale, ed_img, ed_nnscale, ed_bscale, mr_img, wmr_img, mdr_img, ROI, filename,counter=counter) # set figszie

                memfile = BytesIO() 
                f.savefig(memfile)
                document.add_picture(memfile, width=Inches(6)) # set document size
                memfile.close()
                f.clf()
    document.add_section()
    document.add_heading("Podsumowanie i wnioski",1)
    document.add_paragraph("Metoda najblizszego sasiada i interpolacja biliniowa poprawnie wizualizują obraz przy powiekszaniu."
                           "Przy większym powiekszęniu wystepuje wygładzanie krawędzi. "
                           "Metoda NN lepiej zachowuje ostrość  "
                           "-----------------------------"
                           "Jeśli chodzi o testowanie algorytmów zmiejszania  przy silnym pomienszaniu obrazu "
                           "Metoda najblizszego sąsiada i biliniowa tworzą pikselowy efekt. "
                           "Reszta algorytmów radzi sobie lepiej wygładzają obraz co powoduje brak pikselowego efektu")
    document.save(OutputRaportFile) 
